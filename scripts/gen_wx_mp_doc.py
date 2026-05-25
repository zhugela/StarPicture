# -*- coding: utf-8 -*-
"""生成 StarPicture 微信公众号配置说明 Word 文档"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_shading(cell, fill: str):
    try:
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    except Exception:
        pass


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], 'D9E8FF')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def main():
    out = Path(__file__).resolve().parent.parent / 'docs' / '公众号配置说明.docx'
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 封面
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('StarPicture 内娱图库\n微信公众号集成配置说明')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f'文档版本：V1.0    编写日期：{date.today().isoformat()}')
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        '本文档说明 StarPicture 项目微信公众号（自定义菜单、自动回复）与 Spring Boot 后端的对接方式。'
        '公众号与微信小程序共用同一后端服务（端口 8123，上下文路径 /api）。'
    )
    intro.paragraph_format.first_line_indent = Cm(0.74)

    # 1 架构
    doc.add_heading('一、系统架构说明', level=1)
    doc.add_paragraph('整体数据流如下：')
    flow = [
        '用户关注公众号 / 发送消息 → 微信服务器 → 本项目回调接口 /api/wx/mp/portal',
        '后端 Java 服务处理消息（自动回复、菜单点击）→ 返回 XML 文本回复',
        '用户点击菜单「打开图库」→ 跳转微信小程序（同一后端用户与图库 API）',
        '小程序通过 wx.login 调用 /api/user/wx/login，与 Web 端共用数据库与业务接口',
    ]
    for item in flow:
        doc.add_paragraph(item, style='List Bullet')

    add_table(doc, ['组件', '说明', '本项目配置'], [
        ['Spring Boot 后端', '统一 API 服务', 'http://localhost:8123/api'],
        ['公众号 AppID', '公众平台身份标识', 'wx405b15cca532e485（示例，以实际为准）'],
        ['小程序 AppID', '菜单跳转目标', 'wxcc006e978973a40b'],
        ['服务器 Token', 'URL 验签密钥', 'starpicture（与 yml 一致）'],
        ['消息加解密', '当前实现', '明文模式'],
    ], [4, 6, 5])

    # 2 后端配置
    doc.add_heading('二、后端配置文件说明', level=1)
    doc.add_paragraph('配置文件路径：src/main/resources/application-local.yml')
    doc.add_paragraph('关键配置项（wx.mp 为公众号，wx.miniapp 为小程序，不可混用 AppID）：')

    add_table(doc, ['配置项', '含义', '示例值'], [
        ['wx.mp.enabled', '是否启用公众号回调', 'true'],
        ['wx.mp.app-id', '公众号 AppID', 'wx405b15cca532e485'],
        ['wx.mp.app-secret', '公众号 AppSecret', '在公众平台「基本配置」获取，勿泄露'],
        ['wx.mp.token', '服务器配置 Token', 'starpicture'],
        ['wx.miniapp.app-id', '小程序 AppID（菜单跳转）', 'wxcc006e978973a40b'],
    ], [4.5, 5, 5.5])

    doc.add_paragraph('YAML 配置示例：')
    code = doc.add_paragraph()
    code_run = code.add_run('''wx:
  miniapp:
    app-id: wxcc006e978973a40b
    app-secret: <小程序AppSecret>
  mp:
    enabled: true
    app-id: wx405b15cca532e485
    app-secret: <公众号AppSecret>
    token: starpicture''')
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    # 3 接口
    doc.add_heading('三、后端接口清单', level=1)
    add_table(doc, ['接口', '方法', '说明', '鉴权'], [
        ['/api/wx/mp/portal', 'GET', '微信服务器 URL 验证（返回 echostr）', '签名校验，无需登录'],
        ['/api/wx/mp/portal', 'POST', '接收用户消息/事件，自动回复', '签名校验，无需登录'],
        ['/api/wx/mp/menu/create', 'POST', '创建默认自定义菜单', '管理员 JWT'],
        ['/api/user/wx/login', 'POST', '小程序登录', 'code 换 token'],
        ['/api/picture/list/page/vo', 'POST', '公共图库列表', '可选登录'],
    ], [4.5, 1.5, 5.5, 3])

    # 4 公众平台配置
    doc.add_heading('四、微信公众平台服务器配置', level=1)
    doc.add_heading('4.1 配置入口', level=2)
    doc.add_paragraph('正式公众号：登录 https://mp.weixin.qq.com → 设置与开发 → 基本配置 → 服务器配置 → 修改配置。')
    doc.add_paragraph('本地调试可使用「微信公众平台测试号」：https://mp.weixin.qq.com/debug/cgi-bin/sandbox → 接口配置信息。')

    doc.add_heading('4.2 服务器参数填写', level=2)
    add_table(doc, ['字段', '填写内容', '说明'], [
        ['URL（服务器地址）', 'https://<公网域名>/api/wx/mp/portal', '须 HTTPS；本地开发用 cpolar/ngrok 穿透 8123 端口'],
        ['Token（令牌）', 'starpicture', '必须与 application-local.yml 中 wx.mp.token 完全一致'],
        ['EncodingAESKey', '随机生成', '明文模式可生成后保留，不影响当前实现'],
        ['消息加解密方式', '明文模式', '本项目当前为明文 XML 处理'],
    ], [3.5, 5.5, 6])

    doc.add_heading('4.3 本地内网穿透（开发环境）', level=2)
    steps = [
        '启动 Spring Boot，确认 http://127.0.0.1:8123/api/doc.html 可访问',
        '运行 cpolar：cpolar http 8123，获得 HTTPS 地址，如 https://abc123.cpolar.top',
        '完整回调 URL：https://abc123.cpolar.top/api/wx/mp/portal',
        '先启动穿透与后端，再在微信后台点击「提交」',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('4.4 IP 白名单', level=2)
    doc.add_paragraph(
        '调用微信菜单 API 时，需在本机公网 IP 加入公众平台「基本配置 → IP 白名单」，'
        '否则 menu/create 报错 invalid ip not in whitelist。'
    )

    # 5 自动回复
    doc.add_heading('五、自动回复规则', level=1)
    doc.add_paragraph('启用服务器配置后，自动回复由后端代码处理（非公众平台后台手动配置）。')
    add_table(doc, ['触发条件', '回复内容概要'], [
        ['用户关注公众号（subscribe）', '欢迎语：欢迎关注内娱图库，提示回复图库/帮助或点击菜单'],
        ['发送关键词「图库」', '引导点击菜单进入小程序浏览公开图片'],
        ['发送关键词「帮助」', '内娱图库使用说明（菜单、上传等）'],
        ['发送关键词「上传」', '引导打开小程序上传页'],
        ['其他文本消息', '默认回复文案'],
        ['菜单点击 HELP', '同「帮助」回复'],
        ['菜单点击 GALLERY', '同「图库」回复'],
    ], [5, 10])

    # 6 菜单
    doc.add_heading('六、自定义菜单设置', level=1)
    doc.add_heading('6.1 菜单结构', level=2)
    add_table(doc, ['菜单名称', '类型', '行为'], [
        ['打开图库', 'miniprogram（小程序）', '跳转小程序 pages/index/index'],
        ['功能 → 使用帮助', 'click', 'EventKey=HELP，返回帮助文本'],
        ['功能 → 图库说明', 'click', 'EventKey=GALLERY，返回图库说明'],
        ['功能 → 上传图片', 'miniprogram', '跳转小程序 pages/upload/upload'],
    ], [4, 4, 7.5])

    doc.add_heading('6.2 创建菜单步骤', level=2)
    menu_steps = [
        '确保 wx.mp.app-id、app-secret 已配置，且 IP 白名单已添加',
        '注册账号并将 userRole 设为 admin（MySQL：UPDATE user SET userRole=\'admin\' WHERE ...）',
        'POST /api/user/login 获取 JWT Token',
        '调用 POST /api/wx/mp/menu/create，Header：Authorization: Bearer <token>',
        '或在 Knife4j（http://127.0.0.1:8123/api/doc.html）→ 微信公众号 → 创建默认自定义菜单',
        '手机取消关注再关注，或等待数分钟刷新菜单',
    ]
    for s in menu_steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('6.3 关联小程序', level=2)
    doc.add_paragraph(
        '公众号后台需关联小程序 AppID（wxcc006e978973a40b），且与公众号同一微信认证主体，'
        '否则「打开图库」等小程序类型菜单无法正常使用。'
        '路径：广告与服务 → 小程序管理 → 添加，或 设置与开发 → 公众号设置 → 相关小程序。'
    )

    # 7 验收
    doc.add_heading('七、联调验收清单', level=1)
    checklist = [
        'application-local.yml 已配置 wx.mp 三项（app-id、app-secret、token）',
        '内网穿透 URL 提交成功，GET 验签通过',
        '关注公众号收到欢迎语',
        '发送「帮助」收到自动回复',
        'POST /api/wx/mp/menu/create 返回 code=0',
        '手机底部显示自定义菜单',
        '点击「打开图库」可进入小程序',
    ]
    for c in checklist:
        p = doc.add_paragraph(c, style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)

    # 8 故障
    doc.add_heading('八、常见问题排查', level=1)
    add_table(doc, ['现象', '可能原因', '处理方法'], [
        ['Token 验证失败', 'Token 不一致', '统一为 starpicture，重启后端'],
        ['请求超时', '穿透断开或后端未启动', '重启 cpolar 与 Spring Boot'],
        ['404', 'URL 缺少 /api', '使用 .../api/wx/mp/portal'],
        ['invalid ip whitelist', '公网 IP 未加白名单', '基本配置中添加当前 IP'],
        ['菜单无小程序入口', '未关联小程序', '公众号关联 wxcc006e978973a40b'],
        ['48001 创建菜单失败', '订阅号未认证', '使用服务号/认证号或测试号'],
    ], [3.5, 4.5, 7])

    # 9 说明
    doc.add_heading('九、实验报告说明建议', level=1)
    doc.add_paragraph(
        '本项目采用「服务端模式」集成公众号：消息与事件由 /api/wx/mp/portal 统一处理，'
        '自定义菜单通过后端调用微信 cgi-bin/menu/create 接口创建；'
        '公众号负责引流与自动回复，小程序与 Web 端共用同一 Spring Boot 后端与 MySQL 数据库，'
        '实现多前端形态下的业务一体化。'
    )

    doc.add_paragraph()
    foot = doc.add_paragraph('—— 文档结束 ——')
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out))
    print(f'Generated: {out}')


if __name__ == '__main__':
    main()
