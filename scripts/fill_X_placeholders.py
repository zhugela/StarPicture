"""
把 4 份报告里【统计】行的 X 占位符替换成实际数字
"""
import subprocess
from pathlib import Path
from docx import Document

BASE = Path("D:/code/StarPicture/docs/test")

DETAILS = {
    "朱远亮": {"高": 0, "中": 0, "低": 1, "建": 2},
    "李冠燃": {"高": 0, "中": 1, "低": 1, "建": 3},
    "李坤纬": {"高": 0, "中": 0, "低": 1, "建": 4},
    "林景彬": {"高": 0, "中": 2, "低": 1, "建": 3},
}

PDF_NAMES = {
    "朱远亮": "用户模块-安全报告.pdf",
    "李冠燃": "图片模块-安全报告.pdf",
    "李坤纬": "空间模块-安全报告.pdf",
    "林景彬": "文件_公众号模块-安全报告.pdf",
}

DOCX_FILES = {
    "朱远亮": "用户模块-安全报告模板.docx",
    "李冠燃": "图片模块-安全报告模板.docx",
    "李坤纬": "空间模块-安全报告模板.docx",
    "林景彬": "文件_公众号模块-安全报告模板.docx",
}

for member in ["朱远亮", "李冠燃", "李坤纬", "林景彬"]:
    docx_path = BASE / f"{member}_脚本与截图/安全测试/{DOCX_FILES[member]}"
    doc = Document(docx_path)
    d = DETAILS[member]
    target_text = f"【统计】：高危 {d['高']} 个，中危 {d['中']} 个，低危 {d['低']} 个，建议 {d['建']} 个"

    for para in doc.paragraphs:
        if "高危 X 个" in para.text:
            new_text = para.text.replace(
                "高危 X 个，中危 X 个，低危 X 个，建议 X 个",
                target_text.replace("【统计】：", "")
            )
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(new_text)
    doc.save(docx_path)
    print(f"已修: {member}")

# 重新转 PDF
print("\n重新生成 PDF...")
SOFFICE = "C:/Program Files/LibreOffice/program/soffice.exe"
for member in ["朱远亮", "李冠燃", "李坤纬", "林景彬"]:
    template_dir = BASE / f"{member}_脚本与截图/安全测试/"
    for old_pdf in template_dir.glob("*.pdf"):
        old_pdf.unlink()
    docx_path = template_dir / DOCX_FILES[member]
    subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(template_dir)],
        capture_output=True, text=True, timeout=60
    )
    new_pdf = docx_path.with_suffix(".pdf")
    if new_pdf.exists():
        target = template_dir / PDF_NAMES[member]
        new_pdf.rename(target)
        print(f"  → {PDF_NAMES[member]}")

# 验证
print("\n=== 验证 ===")
for member in ["朱远亮", "李冠燃", "李坤纬", "林景彬"]:
    docx_path = BASE / f"{member}_脚本与截图/安全测试/{DOCX_FILES[member]}"
    doc = Document(docx_path)
    for para in doc.paragraphs:
        if "【统计】" in para.text:
            print(f"  [{member}]: {para.text[:120]}")
            break