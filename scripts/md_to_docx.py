"""
把 3 份 md 文档转成 docx
支持：标题、表格、列表、代码块、强调
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = qn(f'w:{edge}')
            elem = tcBorders.find(tag)
            if elem is None:
                elem = OxmlElement(f'w:{edge}')
                tcBorders.append(elem)
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(kwargs[edge]))
            elem.set(qn('w:color'), '999999')

def parse_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    rpr = style.element.rPr
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), '宋体')
        rfonts.set(qn('w:ascii'), 'Times New Roman')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 分隔线 ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(text, level=min(level, 4))
            for run in h.runs:
                run.font.name = '黑体'
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                rfonts = run._element.rPr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = OxmlElement('w:rFonts')
                    run._element.rPr.append(rfonts)
                rfonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # 表格（连续的 | 行）
        if '|' in line and i + 1 < len(lines) and re.search(r'\|[\s\-:|]+\|', lines[i+1]):
            tbl_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                tbl_lines.append(lines[j].rstrip('\n'))
                j += 1
            # 去掉分隔行（第二行是 ---|---|---）
            tbl_data = []
            for idx, tl in enumerate(tbl_lines):
                cells = [c.strip() for c in tl.strip().strip('|').split('|')]
                if idx == 1 and all(re.match(r'^[\s\-:]+$', c) for c in cells):
                    continue
                tbl_data.append(cells)
            if tbl_data:
                ncol = max(len(r) for r in tbl_data)
                # 补齐
                for r in tbl_data:
                    while len(r) < ncol:
                        r.append('')
                t = doc.add_table(rows=len(tbl_data), cols=ncol)
                t.style = 'Light Grid Accent 1'
                t.autofit = True
                for ri, row in enumerate(tbl_data):
                    for ci, val in enumerate(row):
                        cell = t.rows[ri].cells[ci]
                        # 解析单元格内容（支持 `code`）
                        cell.text = ''
                        p = cell.paragraphs[0]
                        parts = re.split(r'(`[^`]+`)', val)
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part[1:-1])
                                run.font.name = 'Consolas'
                                run.font.size = Pt(10)
                            else:
                                run = p.add_run(part)
                                run.font.size = Pt(10)
                        if ri == 0:
                            for run in p.runs:
                                run.bold = True
                            # 蓝色背景
                            tcPr = cell._tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), '305496')
                            shd.set(qn('w:val'), 'clear')
                            tcPr.append(shd)
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                i = j
                continue

        # 列表
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            if i < len(lines):
                i += 1
            for cl in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(0.5)
                # 浅灰背景
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F2F2F2')
                shd.set(qn('w:val'), 'clear')
                pPr.append(shd)
            continue

        # 引用
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph(text, style='Intense Quote')
            i += 1
            continue

        # 普通段落
        text = line
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                run = p.add_run(part)
        i += 1

    doc.save(docx_path)
    print(f"已生成: {docx_path}")

# 转 3 份
pairs = [
    (r"D:/code/StarPicture/docs/test/01_软件测试计划_CS001.md", r"D:/code/StarPicture/docs/test/01_软件测试计划_CS001.docx"),
    (r"D:/code/StarPicture/docs/test/02_软件测试说明_CS002.md", r"D:/code/StarPicture/docs/test/02_软件测试说明_CS002.docx"),
    (r"D:/code/StarPicture/docs/test/03_软件测试报告_CS003.md", r"D:/code/StarPicture/docs/test/03_软件测试报告_CS003.docx"),
]
for src, dst in pairs:
    parse_md_to_docx(src, dst)
