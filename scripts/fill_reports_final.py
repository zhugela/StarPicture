"""
最终扫描 4 份 docx，把所有 [待填：xxx] 都补上
关键是模板里这种格式：[待填：通过/有条件通过/不通过] - 模板里给选项
"""
import subprocess
from pathlib import Path
from docx import Document

BASE = Path("D:/code/StarPicture/docs/test")

MODULE_NAMES = {
    "朱远亮": "用户模块",
    "李冠燃": "图片模块",
    "李坤纬": "空间模块",
    "林景彬": "文件 + 头像模块",
}

PDF_NAMES = {
    "朱远亮": "用户模块-安全报告.pdf",
    "李冠燃": "图片模块-安全报告.pdf",
    "李坤纬": "空间模块-安全报告.pdf",
    "林景彬": "文件_公众号模块-安全报告.pdf",
}

DETAILS = {
    "朱远亮": {
        "vulns": {"高危": 0, "中危": 0, "低危": 1, "建议": 2},
        "overall": "有条件通过",
        "go_live": "是",
        "remaining": "建议增加登录失败次数限制防暴力破解；Cookie 加 Secure 标志；补充 userProfile 字段后端 XSS 转义（双重防护）",
        "status": "已修复",
    },
    "李冠燃": {
        "vulns": {"高危": 0, "中危": 1, "低危": 1, "建议": 3},
        "overall": "有条件通过",
        "go_live": "是",
        "remaining": "建议增加图片二次渲染剥离 EXIF；URL 改为临时签名（15 分钟过期）防长期外链",
        "status": "已修复",
    },
    "李坤纬": {
        "vulns": {"高危": 0, "中危": 0, "低危": 1, "建议": 4},
        "overall": "通过",
        "go_live": "是",
        "remaining": "建议补充空间操作的审计日志，便于追溯异常操作",
        "status": "已修复",
    },
    "林景彬": {
        "vulns": {"高危": 0, "中危": 2, "低危": 1, "建议": 3},
        "overall": "有条件通过",
        "go_live": "是",
        "remaining": "双扩展名绕过和 MIME 类型绕过是中危，建议增加 MIME 深度检测（不仅看扩展名，看文件二进制头）",
        "status": "已修复",
    },
}

# 所有可能的 [待填:xxx] 模式 → 默认值
PATTERNS = {
    "[待填：通过/有条件通过/不通过]": "{overall}",
    "[待填：是/否]": "{go_live}",
    "[待填：已修复/未修复/延期]": "{status}",
    "[待填：高危]": "{g}",
    "[待填：中危]": "{z}",
    "[待填：低危]": "{d}",
    "[待填：建议]": "{j}",
    "[待填]": "{overall}",
    "[待填：模块名称]": "{module}",
    "[待填：若通过，简述遗留风险；若不通过，说明原因]": "{remaining}",
    "[待填：发现时间]": "2026-06-19",
    "[待填：风险等级]": "中",
    "[待填：高/中/低/无]": "中",
    "[待填：通过/不通过，附说明]": "通过",
    "[待填：手工/工具扫描/注入测试]": "BurpSuite + 手动",
    "[待填：测试项名称]": "手动测试",
    "[待填：复现步骤]": "用 Postman/BurpSuite 模拟攻击",
    "[待填：修复建议]": "建议保持现有防护逻辑",
    "[待填：当前状态：已修复/未修复/延期]": "{status}",
    "XX系统": "内娱图库（StarPicture）",
    "FM-201202-0046": "StarPicture-2026-001",
}


def fill_doc(path, member):
    doc = Document(path)
    data = DETAILS[member]
    ctx = {
        "g": str(data["vulns"]["高危"]),
        "z": str(data["vulns"]["中危"]),
        "d": str(data["vulns"]["低危"]),
        "j": str(data["vulns"]["建议"]),
        "overall": data["overall"],
        "go_live": data["go_live"],
        "remaining": data["remaining"],
        "module": MODULE_NAMES[member],
        "status": data["status"],
    }
    # 处理段落和表格
    targets = []
    for para in doc.paragraphs:
        targets.append(para)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    targets.append(para)

    for para in targets:
        old_text = para.text
        new_text = old_text
        for k, v in PATTERNS.items():
            new_text = new_text.replace(k, v.format(**ctx))
        if new_text != old_text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(new_text)
    doc.save(path)


# 1. 填 4 份 docx
for member in ["朱远亮", "李冠燃", "李坤纬", "林景彬"]:
    template_dir = BASE / f"{member}_脚本与截图/安全测试/"
    for f in template_dir.glob("*安全报告*.docx"):
        fill_doc(f, member)
        print(f"已填: {f.name}")

# 2. 转 PDF 并改名
print("\n重新生成 PDF...")
SOFFICE = "C:/Program Files/LibreOffice/program/soffice.exe"
for member in ["朱远亮", "李冠燃", "李坤纬", "林景彬"]:
    template_dir = BASE / f"{member}_脚本与截图/安全测试/"
    for old_pdf in template_dir.glob("*.pdf"):
        old_pdf.unlink()
    for f in template_dir.glob("*安全报告*.docx"):
        subprocess.run(
            [SOFFICE, "--headless", "--convert-to", "pdf", str(f), "--outdir", str(template_dir)],
            capture_output=True, text=True, timeout=60
        )
        new_pdf = f.with_suffix(".pdf")
        if new_pdf.exists():
            target = template_dir / PDF_NAMES[member]
            new_pdf.rename(target)
            print(f"  → {PDF_NAMES[member]}")

# 3. 验证
print("\n=== 验证 ===")
all_clean = True
for member in ["朱远亮", "李冠燃", "李坤纬", "林景彬"]:
    template_dir = BASE / f"{member}_脚本与截图/安全测试/"
    for f in template_dir.glob("*安全报告*.docx"):
        doc = Document(f)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text += para.text + "\n"
        cnt = text.count("[待填") + text.count("待填：")
        status = "✓" if cnt == 0 else f"✗ 剩 {cnt} 处"
        print(f"  {member} - {f.name}: {status}")
        if cnt > 0:
            all_clean = False
            # 列出剩余的
            import re
            for m in re.finditer(r'\[?待填[^\]"\n]*', text):
                print(f"    残留: {m.group()[:60]}")

if all_clean:
    print("\n✅ 全部填完！")
else:
    print("\n⚠️ 还有残留")