# -*- coding: utf-8 -*-
"""生成课程学习总结 Word：2310820053_朱远亮_课程学习总结.docx"""
from datetime import date
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))
from team_members import MEMBER_ROWS

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

FONT_SIZE = Pt(10.5)
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
LINE_SPACING = 1.25


def set_run_font(run, bold=False, size=None):
    run.font.size = size or FONT_SIZE
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    run.bold = bold


def add_title(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, bold=True, size=Pt(size))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, bold=True, size=Pt(14 if level == 1 else 12))


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_table(doc, caption, headers, rows):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            c.text = str(val)
            for para in c.paragraphs:
                for run in para.runs:
                    set_run_font(run)
    doc.add_paragraph()


def main():
    out_dir = Path(__file__).resolve().parent.parent / 'docs'
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / '2310820053_朱远亮_课程学习总结.docx'

    doc = Document()
    doc.styles['Normal'].font.name = FONT_EN
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    doc.styles['Normal'].font.size = FONT_SIZE

    add_title(doc, '课程学习总结', 18)
    add_title(doc, '《Web 开发技术》/ 软件工程课程设计', 14)
    info_lines = [
        '学    号：2310820053',
        '姓    名：朱远亮',
        '学    院：工学院',
        '专    业：计算机科学与技术',
        f'撰写日期：{date.today().strftime("%Y年%m月%d日")}',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = LINE_SPACING
        r = p.add_run(line)
        set_run_font(r)
    doc.add_paragraph()

    add_body(
        doc,
        '本总结围绕本人在本课程中的学习过程与实践项目「星图集（StarPicture）内娱图库管理系统」撰写。'
        '该项目采用 Web 网站、微信小程序与微信公众号多端协同、Spring Boot 统一后端的架构，'
        '覆盖了课堂讲授的 HTML5、CSS3、Ajax、JSON、Servlet 体系以及微信生态开发等核心内容。'
        '以下从主要技术、个人分工、技术困难、收获体会、不足与展望、课程思政六个方面进行系统回顾。'
    )

    # ========== 1 主要技术 ==========
    add_heading(doc, '一、本课程中学习并实践的主要技术', 1)

    add_body(
        doc,
        '（一）HTML5 与 CSS3。在 Web 前端开发中，我使用 Vue 3 单页应用组织页面结构，'
        '底层仍遵循 HTML5 语义化与 CSS3 布局思想。例如 GlobalHeader 导航栏、HomePage 图片卡片网格、'
        'UserLoginPage 与 UserProfilePage 表单页面，均采用 Flex 布局、圆角卡片、响应式间距与 CSS 变量实现统一视觉风格。'
        '小程序端对应使用 WXML 描述结构与 WXSS 完成样式，首页 Hero 搜索区、分类标签 Chip、底部 TabBar 等均体现了'
        '对 HTML5/CSS3 盒模型、定位与媒体查询等基础知识的迁移应用。'
    )
    add_body(
        doc,
        '（二）Ajax 与 JSON。前后端分离架构下，数据交互全部通过 Ajax 异步完成。Web 端基于 Axios 封装 request 模块，'
        '以 JSON 作为请求与响应体格式，调用 /api/user/login、/api/picture/list/page/vo 等 REST 接口；'
        '小程序端使用 wx.request 与 utils/api.js 统一处理 code、message、data 三段式返回。'
        '我实践了 GET/POST 请求、分页参数传递、Bearer Token 与 Cookie Session 双通道认证，'
        '理解了 Ajax 相对传统表单提交在用户体验与接口解耦方面的优势。'
    )
    add_body(
        doc,
        '（三）Servlet 与 Java Web 后端。课程中的 Servlet 知识在 Spring Boot 中以 DispatcherServlet + Controller 形式延续。'
        '我编写的 UserController、PictureController、FileController、WxMpController 等类，'
        '本质上承担原 Servlet 的请求分发、参数解析与响应输出职责，并通过 @RestController 自动将对象序列化为 JSON。'
        '同时学习了 Filter/AOP 思想：AuthInterceptor 切面配合 @AuthCheck 注解实现权限控制，'
        'GlobalExceptionHandler 统一处理 BusinessException，这与 Servlet 时代的过滤器链、异常处理一脉相承。'
    )
    add_body(
        doc,
        '（四）JSON 数据处理。项目中 JSON 不仅用于接口报文，还用于数据库字段存储：picture 表的 tags、urls 等字段以 JSON 字符串或 JSON 类型保存；'
        'Hutool、Jackson 与 MyBatis-Plus 协同完成对象与 JSON 的转换。我在联调中处理了 JavaScript 大整数精度问题，'
        '约定图片 id 以字符串传输，避免前端 Number 失真，加深了对 JSON 作为跨语言数据交换格式的认识。'
    )
    add_body(
        doc,
        '（五）微信公众号 API。我完成了公众号服务器配置（URL 验证、Token 签名校验）、消息加解密接收、'
        '关键词自动回复与自定义菜单创建接口对接。WxMpController 提供 GET/POST /wx/mp/portal，'
        'WxMpService 负责 access_token 缓存与菜单 JSON 提交。实践中理解了公众平台回调机制、'
        '内网穿透在本地调试中的作用，以及启用服务器配置后后台静态菜单失效、需由后端接管的原因。'
    )
    add_body(
        doc,
        '（六）微信小程序开发。我独立完成了小程序 pages/index（公共图库）、upload（上传）、mine（我的）、'
        'space（空间）、profile（个人资料）、color-search（按色搜图）等页面，以及 utils/auth.js 微信登录、'
        'JWT 持久化、custom-tab-bar 自定义导航。掌握了 wx.login、wx.request、wx.uploadFile、'
        'wx.chooseMedia 等 API，理解了小程序生命周期、本地 Storage 与后端 JWT 的配合方式。'
    )
    add_body(
        doc,
        '（七）其他延伸技术。此外还实践了 MySQL 数据库设计、MyBatis-Plus ORM、腾讯云 COS 对象存储、'
        'Knife4j 接口文档、Caffeine 本地缓存、MD5 密码摘要与 JWT 签发校验等，'
        '形成了从表现层到数据层、从同步 HTTP 到第三方云服务的完整 Web 开发知识链。'
    )

    # ========== 2 角色与任务 ==========
    add_heading(doc, '二、小组项目中的具体角色与个人贡献', 1)
    add_body(
        doc,
        '本课程设计小组以「星图集」为共同题目，我（朱远亮，2310820053）担任组长与核心开发，'
        '与小组项目总结报告中的分工描述保持一致。小组成员：林景彬（Web 前端）、李冠燃（测试与文档）、李坤纬（管理端与 UI）。'
        '我主要负责后端业务模块、微信小程序、微信公众号接入及三端联调。'
    )
    add_table(
        doc,
        '表1  星图集小组全体成员分工（摘要）',
        ['学号', '姓名', '主要工作'],
        [(r[0], r[1], r[3][:40] + '…') for r in MEMBER_ROWS],
    )
    add_table(
        doc,
        '表2  本人（朱远亮）具体贡献',
        ['模块', '具体任务', '完成情况'],
        [
            ('用户模块', '注册/登录/微信登录/JWT、个人资料与头像上传接口', '已完成'),
            ('图片模块', '上传、分页检索、审核、公共库 nullSpaceId 策略', '已完成'),
            ('空间模块', '空间 CRUD、配额级别、空间内图片管理', '已完成'),
            ('Web 前端', '导航重构、登录与个人资料页、Axios Token 拦截', '已完成'),
            ('微信小程序', 'TabBar、首页筛选、上传/我的/空间/资料页', '已完成'),
            ('微信公众号', 'portal 回调、自动回复、菜单 create 接口', '已完成'),
            ('测试与文档', '三端联调、缺陷修复、配置说明与课程报告', '已完成'),
        ],
    )
    add_body(
        doc,
        '在协作方式上，我负责制定前后端接口约定（统一 BaseResponse、错误码 40100/40101），'
        '并推动小组成员按 Controller-Service-Mapper 分层提交代码；'
        '对于 UI 与交互细节，我与负责演示与测试的同学确认用例后再合并。'
        '个人代码量主要集中在 backend 的 user/picture/space/wx 包、star-picture-vue 的 pages 与 components、'
        '以及 mini 小程序 utils 与 pages 目录，占小组可运行版本的主体部分。'
    )

    # ========== 3 技术困难 ==========
    add_heading(doc, '三、项目开发中的典型技术困难及解决方法', 1)

    add_body(
        doc,
        '困难一：微信小程序公共图库首页长期显示「暂无图片」，而后端数据库中已有数据。'
        '排查发现 list 接口在未传 spaceId 时需设置 nullSpaceId=true 才能查询公共空间；'
        '小程序首页 buildQuery 最初遗漏该字段，导致查询条件与预期不符。'
        '解决方法是在 index.js 的 buildQuery 中显式加入 nullSpaceId: true，'
        '并对照后端 PictureController.preparePictureQueryForList 逻辑确认 reviewStatus 仅展示审核通过内容。'
        '该问题让我认识到前后端对「公共图库」语义必须文档化，不能依赖隐式默认。'
    )
    add_body(
        doc,
        '困难二：图片 multipart 上传时 spaceId、picName 等参数后端接收为空。'
        'FileController 原先未对 PictureUploadRequest 使用 @ModelAttribute，'
        'Spring MVC 在 multipart 请求中无法自动绑定表单字段。'
        '解决方法是增加 @ModelAttribute 注解并在 request 为 null 时 new 默认对象；'
        '小程序 upload.js 则在 ensureLogin 后再调用 wx.uploadFile，保证 Authorization 头携带 JWT。'
        '由此理解了文件上传与普通 JSON POST 在参数绑定上的差异。'
    )
    add_body(
        doc,
        '困难三：微信公众号本地联调与菜单创建失败。'
        '启用服务器配置后，需公网 URL 供微信服务器回调，本地使用 cpolar 内网穿透；'
        '调用菜单创建 API 时报 IP 白名单错误，需在公众平台配置出口公网 IP。'
        '此外 application.yml 中 keyword-reply 多行 Map 写法曾导致 Spring Boot 启动失败，'
        '最终改为 Java 代码内默认常量。'
        '这一系列问题让我体会到第三方平台集成除编码外，还 heavily 依赖运维配置与环境准备。'
    )
    add_body(
        doc,
        '困难四（补充）：三端登录态不一致。Web 端最初仅依赖 Session Cookie，'
        '小程序使用 JWT，导致部分接口在 Web 刷新后偶发未登录；'
        '我在 request.ts 增加 localStorage 读取 sp_token 并写入 Authorization 头，'
        '登录页保存后端返回的 token，退出时 clearToken，与小程序 auth.js 策略对齐。'
        '同时在 api.js 的 ok 封装中对 code=40100 执行 clearLogin，避免过期 token 反复请求。'
    )

    # ========== 4 个人收获 ==========
    add_heading(doc, '四、个人收获', 1)
    add_body(
        doc,
        '技术能力方面，我从单一写页面或写接口，升级为能独立打通「数据库—后端—前端—小程序—第三方 API」链路。'
        '特别是 REST 设计、异常码规范、VO 脱敏、逻辑删除与审核状态机等内容，'
        '使我对企业级 Web 项目的常见模式有了直观认识。调试能力显著提升：'
        '会使用 Knife4j、浏览器 Network、微信开发者工具 Console 交叉定位问题。'
    )
    add_body(
        doc,
        '团队协作方面，我体会到接口先行、约定优于配置的重要性。'
        '早期若未统一分页参数与 id 类型，后期合并成本很高。'
        '我在小组中主动编写联调步骤与配置说明，减少其他同学环境搭建时间，'
        '也锻炼了技术表达与文档写作能力。'
    )
    add_body(
        doc,
        '工程思维方面，我学会了分层与单一职责：Controller 只做校验与转发，'
        '业务规则放在 Service，权限用切面集中处理。'
        '对「能跑」与「可维护、可扩展、安全」的差距有了更清晰认知，'
        '例如公共图库必须过滤未审核图片、管理员接口不能返回密码字段等，'
        '这些都是工程化细节，而非单纯功能堆砌。'
    )

    # ========== 5 不足与方向 ==========
    add_heading(doc, '五、不足之处与今后努力方向', 1)
    add_body(
        doc,
        '不足之处：其一，自动化测试不足，主要依赖手工用例，单元测试与接口测试覆盖率低；'
        '其二，性能优化停留在 Caffeine 本地缓存，未引入 Redis 与 CDN，大图列表高并发场景未经压测；'
        '其三，部分历史代码规范不统一，个别 Controller 参数校验与错误码仍有改进空间；'
        '其四，对 Servlet 源码级原理、Spring Security 体系的学习仍不够深入，更多停留在应用层。'
    )
    add_body(
        doc,
        '今后努力方向：系统补学 Java Web 底层（Servlet 规范、HTTP 协议、Tomcat 机制）与 Spring Security；'
        '为项目补充 JUnit/MockMvc 测试与 GitHub Actions 持续集成；'
        '探索 Elasticsearch 全文检索提升图库搜索体验；'
        '关注国产云与信创环境部署；继续参与开源或实验室项目，将课程项目迭代为可展示的个人作品集。'
    )

    # ========== 6 课程思政 ==========
    add_heading(doc, '六、课程思政：对「技术赋能社会」「科技报国」与「工匠精神」的理解', 1)

    ideology_paras = [
        '通过本课程与「星图集」项目的实践，我对「技术赋能社会」有了更加具体而非空泛的理解。'
        '技术赋能社会，并非指堆砌花哨的功能，而是让普通人以更低成本获得过去只有专业机构才具备的能力。'
        '我们的图库系统面向内娱爱好者与内容创作者，提供素材上传、分类检索、空间管理与内容审核，'
        '本质上是在用 Web 与移动互联网技术降低信息整理与分享的门槛。'
        '当一位用户可以在小程序中上传作品、在通过审核后进入公共图库被他人检索使用时，'
        '技术就在 silently 地连接个体创造与集体受益。'
        '这让我意识到，计算机专业学生所写的每一行代码，都可能影响真实用户的体验与权益，'
        '因此必须在功能设计之外考虑内容合规、隐私保护与公平访问，'
        '例如公共库仅展示审核通过图片、密码加密与 VO 脱敏等，'
        '正是「赋能」与「守底线」并重的体现。',
        '关于「科技报国」，我的理解是：报国不一定立刻体现为重大工程，'
        '而是把所学转化为解决实际问题的能力，并在关键领域保持自主可控意识。'
        '在本项目中，后端、数据库、对象存储、前端框架均是我们亲手搭建与配置的，'
        '微信公众号与小程序接口对接让我们接触到国产数字生态的核心入口。'
        '当菜单 API 因 IP 白名单、回调因网络环境受阻时，我并未简单放弃，'
        '而是通过内网穿透、阅读官方文档、在后端实现 token 缓存与消息处理等方式逐一攻克。'
        '这种在约束条件下仍坚持完成集成的心态，正是科技工作者面对「卡脖子」环境时需要的韧劲。'
        '未来我希望在数据库、中间件、操作系统等基础软件方向继续深耕，'
        '把「会用框架」提升为「懂原理、能优化、可替代」，'
        '以专业能力服务国家信息化与数字中国建设，这是我对科技报国的朴素承诺。',
        '「工匠精神」在项目开发中体现为对细节的反复打磨与对质量的敬畏。'
        '记得为解决小程序首页空白，我反复对照后端 SQL 条件与前端 query 对象；'
        '为统一三端 UI，多次调整 CSS 变量与小程序 rpx 间距；'
        '为通过 npm run build 与课程报告字数要求，耐心修复 TypeScript 类型错误并完善文档。'
        '工匠精神不是慢，而是不将就：接口返回结构不统一就 refactor，'
        '401 未清 token 就补拦截器，multipart 参数丢失就查绑定注解。'
        '我也认识到，工匠精神包含对他人的负责——'
        '写清楚的 README 与联调文档，与写正确的代码同样重要，'
        '因为软件是社会协作的产物，下游同学与未来的自己都是「用户」。',
        '结合三者，我认为技术赋能社会是目标，科技报国是方向，工匠精神是路径。'
        '作为工学院计算机科学与技术专业的学子，朱远亮这个名字背后是一份具体的责任：'
        '既要在课堂上掌握 HTML5、Ajax、Servlet、JSON 与微信开发等硬技能，'
        '更要在价值观上明确技术为谁服务、为何服务。'
        '「星图集」或许只是课程中的一个项目，但它训练了我从需求到上线的完整工程经验。'
        '今后无论从事 Web 开发、后端架构还是其他计算机方向，'
        '我都将以严谨规范、持续学习、服务社会为准则，'
        '把个人成长融入国家科技进步与数字文明建设之中，'
        '用实际行动诠释新时代青年学子的担当。',
    ]
    for para in ideology_paras:
        add_body(doc, para)

    add_body(
        doc,
        '综上所述，本课程不仅教会我如何开发网站与小程序，更促使我思考技术与社会的关系。'
        '我将以本次项目为起点，继续完善 StarPicture 系统，并在学习与工作中践行技术赋能、科技报国与工匠精神。'
    )

    doc.save(out)
    text = '\n'.join(p.text for p in doc.paragraphs)
    cn = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    # 思政部分字数
    idx = text.find('六、课程思政')
    si = sum(1 for c in text[idx:] if '\u4e00' <= c <= '\u9fff') if idx >= 0 else 0
    print(f'已生成: {out}')
    print(f'总汉字约: {cn}，思政部分汉字约: {si}')


if __name__ == '__main__':
    main()
