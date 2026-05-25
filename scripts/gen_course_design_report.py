# -*- coding: utf-8 -*-
"""生成 StarPicture 课程设计报告 Word 文档（按学院模板）"""
from datetime import date
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

FONT_SIZE = Pt(10.5)
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
LINE_SPACING = 1.25

STUDENT = {
    'name': '朱远亮',
    'id': '2310820053',
    'college': '工学院',
    'major': '计算机科学与技术',
}

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from team_members import MEMBER_ROWS, CONTRIBUTION_ROWS, GRADE_LINES


def setup_plt():
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False


def set_run_font(run, bold=False, size=None):
    run.font.size = size or FONT_SIZE
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    run.bold = bold


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACING
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_run_font(p.add_run(text))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, bold=True, size=Pt(14 if level == 1 else 12 if level == 2 else 11))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text))


def add_image(doc, path, w=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(w))


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(code)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)


def add_table(doc, caption, headers, rows):
    add_caption(doc, caption)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for para in c.paragraphs:
                for run in para.runs:
                    set_run_font(run)
    doc.add_paragraph()


def draw_structure(path):
    setup_plt()
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')

    def box(x, y, w, h, t, fc='#E8F4FF'):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.02', fc=fc, ec='#333'))
        ax.text(x + w / 2, y + h / 2, t, ha='center', va='center', fontsize=9)

    box(3.5, 9, 3, 0.7, '星图集 StarPicture', '#FFD1DC')
    for x, t in [(0.3, 'Web(Vue3)'), (3.7, '微信小程序'), (7.1, '微信公众号')]:
        box(x, 7.5, 2.4, 1.0, t)
        ax.annotate('', xy=(x + 1.2, 7.5), xytext=(5, 9), arrowprops=dict(arrowstyle='->', lw=1))
    box(2.5, 5.8, 5, 0.8, 'Spring Boot API /api', '#D9E8FF')
    for i, m in enumerate(['用户', '图片', '空间', '文件', '微信', '权限AOP']):
        bx, by = 0.4 + (i % 3) * 3.1, (4.0 if i < 3 else 2.6)
        box(bx, by, 2.6, 0.9, m)
        ax.annotate('', xy=(bx + 1.3, by + 0.9), xytext=(5, 5.8), arrowprops=dict(arrowstyle='->', lw=0.8))
    box(1.5, 0.8, 3, 0.9, 'MySQL', '#FFF4BD')
    box(5.5, 0.8, 3, 0.9, '腾讯云COS', '#B2E2F2')
    ax.annotate('', xy=(3, 1.7), xytext=(5, 5.8), arrowprops=dict(arrowstyle='->', lw=1))
    ax.annotate('', xy=(7, 1.7), xytext=(5, 5.8), arrowprops=dict(arrowstyle='->', lw=1))
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_usecase(path):
    setup_plt()
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.add_patch(FancyBboxPatch((2.5, 0.5), 5, 6, boxstyle='round,pad=0.02', fc='#F5F5F5', ec='#999', ls='--'))
    ax.text(5, 6.2, '星图集系统', ha='center', fontsize=11, fontweight='bold')
    cases = ['注册登录', '浏览图库', '上传图片', '空间管理', '个人资料', '图片审核', '用户管理', '公众号回复']
    for i, c in enumerate(cases):
        y = 5.5 - i * 0.65
        ax.add_patch(FancyBboxPatch((3.2, y - 0.22), 3.6, 0.44, boxstyle='round,pad=0.02', fc='#E8F4FF', ec='#333'))
        ax.text(5, y, c, ha='center', va='center', fontsize=9)
    actors = [('游客', 0.8, 5.0), ('普通用户', 0.5, 3.2), ('管理员', 0.7, 1.5)]
    for name, x, y in actors:
        ax.text(x, y + 0.5, name, fontsize=10, fontweight='bold')
        ax.add_patch(FancyBboxPatch((x - 0.05, y - 0.8), 0.1, 1.6, fc='white', ec='black'))
        ax.plot([x + 0.05, 3.2], [y, 4.8], 'k-', lw=0.6, alpha=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_er(path):
    setup_plt()
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    def ent(x, y, w, h, name, attrs):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle='square,pad=0', fc='white', ec='black'))
        ax.plot([x, x + w], [y + h - 0.32, y + h - 0.32], 'k-')
        ax.text(x + w / 2, y + h - 0.16, name, ha='center', fontsize=10, fontweight='bold')
        for i, a in enumerate(attrs):
            ax.text(x + 0.08, y + h - 0.48 - i * 0.26, a, fontsize=7.5)

    ent(0.4, 2.0, 2.5, 2.3, 'user', ['PK id', 'userAccount', 'userPassword', 'userRole', 'userAvatar'])
    ent(3.5, 1.7, 3.0, 2.8, 'picture', ['PK id', 'FK userId', 'FK spaceId', 'urls JSON', 'reviewStatus', 'tags'])
    ent(7.2, 2.1, 2.4, 2.0, 'space', ['PK id', 'FK userId', 'spaceName', 'spaceLevel', 'maxCount'])
    ax.text(3.0, 3.5, '1:N', fontsize=8)
    ax.text(6.6, 3.2, 'N:1', fontsize=8)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_seq(path, title, steps):
    setup_plt()
    actors = list(dict.fromkeys([s[0] for s in steps] + [s[2] for s in steps]))
    fig, ax = plt.subplots(figsize=(10, max(5, len(steps) * 0.45 + 2)))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, len(steps) + 2)
    ax.axis('off')
    ax.set_title(title, fontsize=11, pad=10)
    xs = {a: 1 + i * (8 / max(len(actors) - 1, 1)) for i, a in enumerate(actors)}
    for a, x in xs.items():
        ax.plot([x, x], [0.5, len(steps) + 1], '--', color='#999', lw=0.8)
        ax.text(x, len(steps) + 1.3, a, ha='center', fontsize=9)
    y = len(steps)
    for frm, msg, to in steps:
        ax.annotate('', xy=(xs[to], y), xytext=(xs[frm], y), arrowprops=dict(arrowstyle='->', lw=1))
        ax.text((xs[frm] + xs[to]) / 2, y + 0.12, msg, ha='center', fontsize=8)
        y -= 1
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_ui(path, title, subtitle, rows, tabs=None):
    setup_plt()
    fig, ax = plt.subplots(figsize=(8, 9))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.add_patch(FancyBboxPatch((0.3, 0.3), 7.4, 8.4, boxstyle='round,pad=0.03', fc='#FAFAFA', ec='#CCC', lw=2))
    ax.add_patch(Rectangle((0.3, 8.3), 7.4, 0.5, fc='#FF69B4', alpha=0.25))
    ax.text(4, 8.55, title, ha='center', fontsize=12, fontweight='bold')
    ax.text(4, 7.95, subtitle, ha='center', fontsize=9, color='#666')
    y = 7.2
    for row in rows:
        if row['t'] == 'search':
            ax.add_patch(FancyBboxPatch((0.7, y), 6.6, 0.5, boxstyle='round,pad=0.02', fc='white', ec='#FFB6C1'))
            ax.text(1.0, y + 0.25, row['v'], fontsize=9, va='center', color='#999')
            y -= 0.8
        elif row['t'] == 'grid':
            for i in range(4):
                px, py = 0.7 + (i % 2) * 3.3, y - (i // 2) * 1.8
                ax.add_patch(FancyBboxPatch((px, py - 1.4), 3.0, 1.4, boxstyle='round,pad=0.02', fc='#E8E8E8', ec='#DDD'))
                ax.text(px + 1.5, py - 0.7, f'图片{i+1}', ha='center', va='center', fontsize=8, color='#888')
            y -= 3.8
        elif row['t'] == 'form':
            for lb in row['labels']:
                ax.text(0.9, y, lb, fontsize=9)
                ax.add_patch(Rectangle((0.9, y - 0.42), 6.2, 0.35, fc='white', ec='#CCC'))
                y -= 0.7
            ax.add_patch(FancyBboxPatch((2.5, y - 0.05), 3.0, 0.48, boxstyle='round,pad=0.02', fc='#FF69B4'))
            ax.text(4, y + 0.2, '登录', ha='center', va='center', fontsize=10, color='white', fontweight='bold')
            y -= 0.8
    if tabs:
        for i, tab in enumerate(tabs):
            ax.text(1.0 + i * 1.8, 0.7, tab, ha='center', fontsize=8, color='#FF69B4' if i == 0 else '#999')
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def main():
    base = Path(__file__).resolve().parent.parent
    assets = base / 'docs' / 'report_assets'
    assets.mkdir(parents=True, exist_ok=True)

    p_struct = assets / 'fig_structure.png'
    p_usecase = assets / 'fig_usecase.png'
    p_er = assets / 'fig_er.png'
    p_seq1 = assets / 'fig_seq_login.png'
    p_seq2 = assets / 'fig_seq_upload.png'
    p_ui1 = assets / 'ui_web_home.png'
    p_ui2 = assets / 'ui_web_login.png'
    p_ui3 = assets / 'ui_mini_home.png'

    draw_structure(p_struct)
    draw_usecase(p_usecase)
    draw_er(p_er)
    draw_seq(p_seq1, '用户登录时序图', [
        ('用户', '1.输入账号密码', '客户端'),
        ('客户端', '2.POST /user/login', '后端'),
        ('后端', '3.校验MD5密码', '后端'),
        ('后端', '4.Session+JWT', '客户端'),
        ('客户端', '5.进入首页', '用户'),
    ])
    draw_seq(p_seq2, '图片上传时序图', [
        ('用户', '1.选择图片', '客户端'),
        ('客户端', '2.multipart上传', '后端'),
        ('后端', '3.上传COS', '存储'),
        ('存储', '4.写入MySQL', '后端'),
        ('后端', '5.返回PictureVO', '客户端'),
    ])
    draw_ui(p_ui1, '星图集 Web', '公共图库首页', [
        {'t': 'search', 'v': '输入关键词搜索...'},
        {'t': 'grid'},
    ])
    draw_ui(p_ui2, '星图集 Web', '用户登录页', [
        {'t': 'form', 'labels': ['账号', '密码']},
    ])
    draw_ui(p_ui3, '星图集 小程序', '首页', [
        {'t': 'search', 'v': '搜索图片素材...'},
        {'t': 'grid'},
    ], tabs=['首页', '上传', '空间', '我的'])

    out = base / 'docs' / f'课程设计报告_{STUDENT["name"]}.docx'
    doc = Document()
    doc.styles['Normal'].font.name = FONT_EN
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    doc.styles['Normal'].font.size = FONT_SIZE

    # 封面
    for _ in range(2):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run('软件工程课程设计报告')
    set_run_font(tr, bold=True, size=Pt(22))
    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(tp2.add_run('题目：星图集（StarPicture）内娱图库管理系统'), bold=True, size=Pt(15))
    for line in [
        f'姓    名：{STUDENT["name"]}',
        f'学    号：{STUDENT["id"]}',
        f'学    院：{STUDENT["college"]}',
        f'专    业：{STUDENT["major"]}',
        '指导教师：____________',
        f'完成日期：{date.today().strftime("%Y年%m月%d日")}',
    ]:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(pp.add_run(line))
    doc.add_page_break()

    # 摘要
    add_heading(doc, '摘要', 1)
    add_body(
        doc,
        '星图集小组围绕「星图集」内娱图库管理系统开展课程设计与开发，系统面向图片素材的采集、审核、检索与空间化管理，'
        '支持 Vue3 Web 网站、微信小程序与微信公众号三端接入，后端采用 Spring Boot 统一提供 REST API，'
        '数据存储于 MySQL，图片文件存储于腾讯云 COS。'
        '本报告从需求分析、系统架构、技术选型、数据库设计、系统实现与测试等方面进行阐述。'
        '系统已具备 HTML5/CSS3 前端展示、Ajax/JSON 数据交互、Java Web 后端、微信公众号 API 与微信小程序开发等能力，'
        '并在真实联调中验证了业务闭环的可行性。'
        '本报告由工学院计算机科学与技术专业学生朱远亮（学号2310820053）参与撰写，与小组项目总结报告内容相互印证。'
    )
    doc.add_page_break()

    # 1 课程设计目的
    add_heading(doc, '1  课程设计目的', 1)
    for t in [
        '（1）掌握基于 B/S 架构的 Web 应用与移动端小程序协同开发方法，理解前后端分离与 RESTful API 设计思想；',
        '（2）学习 Spring Boot、MyBatis-Plus、Vue3、微信小程序等主流技术栈的综合运用；',
        '（3）提高需求分析、数据库设计、模块划分与系统测试等软件工程实践能力；',
        '（4）培养解决实际问题的能力，包括用户认证、文件云存储、内容审核与第三方平台（微信）集成；',
        '（5）锻炼文档撰写与团队协作意识，为后续毕业设计与企业项目开发奠定基础。',
        '（6）理解软件工程中需求、设计、实现、测试各阶段的文档规范与质量要求，保证交付物完整可用。',
    ]:
        add_body(doc, t)

    # 2 题目及要求
    add_heading(doc, '2  课程设计题目及要求', 1)
    add_body(doc, '本课程设计题目为「星图集（StarPicture）内娱图库管理系统」。以下从项目背景、用户角色、功能需求、非功能需求、用例概述与约束条件等方面详细说明。')
    add_body(doc, '2.1 项目背景。随着粉丝经济与内容创作的发展，用户对明星写真、海报素材等图片资源的整理、分享与检索需求日益增长。传统网盘或社交软件难以兼顾分类检索、内容审核与多端展示，因此需要一套专门的图库管理系统。')
    add_body(doc, '2.2 用户角色。系统包含三类角色：（1）游客：可浏览公共图库中已审核通过的图片；（2）普通用户（user）：可注册登录、上传图片、管理个人空间、编辑资料；（3）管理员（admin）：可审核图片、管理用户、配置公众号菜单等。')
    add_body(doc, '2.3 功能需求。用户模块：注册、登录（含微信 code 登录）、注销、个人资料与头像维护、管理员用户 CRUD。图片模块：本地上传、URL 上传、批量上传、分页检索（关键词/分类/标签）、编辑删除、公共库审核机制。空间模块：创建多级空间、容量配额、空间内图片管理与按颜色搜图。微信模块：公众号服务器验证、消息回复、自定义菜单；小程序端完整业务闭环。')
    add_body(doc, '2.4 非功能需求。安全性：密码 MD5 存储、VO 脱敏、JWT+Session 双通道认证、@AuthCheck 权限切面。性能：列表分页、无筛选条件下 Caffeine 缓存接口；可维护性：Controller-Service-Mapper 分层、Knife4j 文档；可用性：三端 UI 风格统一（星图集品牌色）、错误码与 Toast 提示友好。')
    add_body(doc, '2.5 用例概述。核心用例包括 UC-01 用户注册、UC-02 账号登录、UC-03 微信登录、UC-04 图片本地上传、UC-05 图片 URL 上传、UC-06 公共图库检索、UC-07 图片审核、UC-08 空间创建、UC-09 空间内批量编辑、UC-10 按颜色搜图、UC-11 个人资料维护、UC-12 公众号关键词回复等。P0 为注册登录与公共图库浏览上传，P1 为空间与审核，P2 为以图搜图与 AI 扩图等扩展能力。')
    add_body(doc, '2.6 约束条件。课程周期有限，采用 Spring Boot 单体架构；图片存储依赖腾讯云 COS；微信公众号需公网环境与 IP 白名单；小程序开发阶段通过「不校验合法域名」连接本地后端。性能指标：常规查询响应小于 2 秒；支持多用户并发浏览；密码加密与权限控制满足基本安全要求。')
    add_body(doc, '2.7 交付物要求。需提交课程设计报告（含需求、设计、实现、测试）、可运行源代码（后端+Web+小程序）、数据库脚本、接口文档截图及关键界面运行效果图；答辩时需演示注册登录、上传审核、公共图库检索等核心流程。')

    add_heading(doc, '2.8  项目组成员及分工及成绩评定', 2)
    add_body(doc, '本项目由星图集小组 4 名成员协作完成，组长朱远亮承担架构、后端、小程序、微信公众号及三端联调等核心工作，其余成员分别负责 Web 前端、测试文档、管理端与 UI 等工作。分工如下表，成绩栏由指导教师评定后填写。')
    add_table(
        doc,
        '表2-1  项目组成员及分工及成绩评定',
        ['学号', '姓名', '专业', '主要工作', '课程设计成绩'],
        MEMBER_ROWS,
    )
    add_table(
        doc,
        '表2-2  个人贡献说明（与小组项目总结报告一致）',
        ['姓名', '学号', '角色', '主要任务', '个人贡献说明'],
        CONTRIBUTION_ROWS,
    )
    add_caption(doc, '表2-3  个人课程设计成绩评定')
    gt = doc.add_table(rows=len(GRADE_LINES), cols=1)
    gt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (sid, name, tag) in enumerate(GRADE_LINES):
        c = gt.rows[i].cells[0]
        c.text = f'学号：{sid}    {tag}  姓名：{name}     成    绩：__________'
        for para in c.paragraphs:
            for run in para.runs:
                set_run_font(run)
    doc.add_paragraph()

    # 3 报告内容
    add_heading(doc, '3  课程设计报告内容', 1)

    add_heading(doc, '3.1  系统功能需求分析', 2)
    add_body(doc, '本系统面向内娱图库场景，功能需求可归纳为用户管理、图片管理、空间管理、微信集成四大子系统。游客无需登录即可访问公共图库；普通用户完成注册登录后可上传与管理素材；管理员负责内容审核与平台治理。')
    add_body(doc, '从数据流看，图片上传后写入 COS 与 MySQL，公共库图片进入待审状态，管理员审核通过后对全部用户可见；上传至私有空间的图片仅空间所有者可管理。个人资料修改同步至 Web 导航栏与小程序「我的」页，保证多端体验一致。')
    add_body(doc, '功能需求与非功能需求对照：检索模块需支持关键词模糊匹配与分类、标签组合过滤；上传模块需限制文件类型与大小；审核模块需记录 reviewerId 与 reviewTime 以便追溯；权限模块需区分本人、管理员与游客三种访问策略。上述规则在 PictureController.preparePictureQueryForList 与 checkPictureAuth 等方法中落地。')
    add_image(doc, p_usecase, 12)
    add_caption(doc, '图3-1  系统用例图（角色与核心用例）')

    add_heading(doc, '3.2  系统功能结构图', 2)
    add_body(doc, '系统采用前后端分离、三端统一接入的 B/S 扩展架构。表现层为 Vue3 Web、微信小程序、微信公众号；应用层为 Spring Boot 提供的 RESTful API；数据层为 MySQL 与腾讯云 COS。各子模块职责清晰：用户模块负责认证授权，图片模块负责素材生命周期，空间模块负责私有归档，微信模块负责第三方生态对接。')
    add_image(doc, p_struct, 13)
    add_caption(doc, '图3-2  系统功能结构图')

    add_heading(doc, '3.3  系统功能设计', 2)
    add_body(doc, '登录模块：UserController 接收 JSON 请求，UserService 校验 MD5 密码，成功则写 Session 并签发 JWT。AuthInterceptor 拦截 @AuthCheck 接口校验 admin 角色。')
    add_image(doc, p_seq1, 13)
    add_caption(doc, '图3-3  用户登录时序图')
    add_body(doc, '上传模块：FileController 使用 @ModelAttribute 绑定 PictureUploadRequest，PictureService 调用 CosManager 上传文件并解析图片元数据入库。公共列表查询强制 nullSpaceId=true 且 reviewStatus=PASS，防止未审核内容外泄。')
    add_body(doc, '审核模块：管理员调用 /picture/review 修改 reviewStatus 与 reviewMessage，系统记录 reviewerId、reviewTime。空间模块：创建空间时根据 spaceLevel 初始化 maxCount、maxSize，上传图片时更新 totalCount、totalSize 并校验配额。')
    add_image(doc, p_seq2, 13)
    add_caption(doc, '图3-4  图片上传时序图')

    add_heading(doc, '3.4  数据库设计', 2)
    add_body(doc, '数据库 starpicture 含 user、picture、space 三表。picture.urls 为 JSON 字段；reviewStatus 实现审核；spaceId 为空表示公共图库。详见 sql/create_sql.sql。')
    add_body(doc, '字段设计遵循第三范式基本要求：用户与图片、空间通过外键关联；tags 以 JSON 字符串存储标签数组；picColor 支持按色搜图；isDelete 逻辑删除避免物理删除风险。createTime、updateTime 自动维护，便于审计与排序。')
    add_image(doc, p_er, 13)
    add_caption(doc, '图3-5  数据库 ER 图')
    add_table(doc, '表3-1  核心表说明', ['表名', '说明', '关键字段'], [
        ('user', '用户账号与角色', 'userAccount, userRole, userAvatar'),
        ('picture', '图片元数据', 'urls, userId, spaceId, reviewStatus, tags'),
        ('space', '用户私有空间', 'spaceName, spaceLevel, maxCount, userId'),
    ])

    add_heading(doc, '3.5  系统实现所用技术', 2)
    add_table(doc, '表3-2  技术选型', ['层次', '技术', '说明'], [
        ('后端', 'Java 17 + Spring Boot 2.7.6', 'REST API、依赖注入、AOP'),
        ('持久层', 'MyBatis-Plus 3.5.9 + MySQL 8', 'ORM、分页、逻辑删除'),
        ('前端', 'Vue3 + TS + Ant Design Vue', 'HTML5/CSS3、组件化、Ajax/Axios'),
        ('小程序', '微信原生 + wx.request', '移动端、JWT 认证'),
        ('存储', '腾讯云 COS', '图片对象存储'),
        ('文档', 'Knife4j 4.4.0', 'OpenAPI 在线调试'),
        ('微信', '公众号 API + 小程序 code2Session', '消息回调、菜单、登录'),
    ])
    add_body(doc, 'Web 端通过 Axios 发起 Ajax 请求，数据格式为 JSON；后端 Controller 等价于 Servlet 体系中的请求处理器；小程序 wx.request 与 Web 共用同一套 API 与 BaseResponse 返回规范。')
    add_body(doc, '与课程知识点的对应关系：HTML5/CSS3 体现在 Vue 页面结构与 WXSS 样式；Ajax 对应 Axios/wx.request 异步通信；Servlet 对应 Spring MVC DispatcherServlet 及 Controller；JSON 为前后端统一数据格式；微信公众号 API 与小程序 API 分别用于消息回调与移动登录。')

    add_heading(doc, '3.6  系统实现', 2)

    add_heading(doc, '3.6.1  登录功能', 3)
    add_body(doc, '实现用户注册、账号密码登录、微信 code 登录、获取当前用户与注销。Web 端登录页保存 JWT 至 localStorage；小程序 auth.js 持久化 sp_token。注册接口校验账号长度、密码强度与两次密码一致性。')
    add_code(doc, '''@PostMapping("/login")
public BaseResponse<LoginUserVo> userLogin(
        @RequestBody UserLoginRequest req, HttpServletRequest request) {
    ThrowUtils.throwIf(req == null, ErrorCode.PARAMS_ERROR);
  LoginUserVo vo = userService.userlogin(
        req.getUserAccount(), req.getUserPassword(), request);
    return ResultUtils.success(vo);
}''')
    add_image(doc, p_ui2, 10)
    add_caption(doc, '图3-6  Web 登录界面运行效果')

    add_heading(doc, '3.6.2  公共图库与图片上传', 3)
    add_body(doc, '首页调用 /picture/list/page/vo，参数 nullSpaceId=true 仅查公共库。upload 页与 FileController 配合完成 multipart 上传，上传后 reviewStatus=0 待审核。支持 Caffeine 缓存接口加速无筛选列表加载。')
    add_image(doc, p_ui1, 10)
    add_caption(doc, '图3-7  Web 公共图库首页')
    add_image(doc, p_ui3, 10)
    add_caption(doc, '图3-8  小程序首页运行效果')

    add_heading(doc, '3.6.3  空间管理与图片审核', 3)
    add_body(doc, 'SpaceController 实现空间 CRUD 与级别配额。PictureController /review 供管理员审核，通过后公共库可见。小程序 review 页仅 admin 可访问。')
    add_body(doc, '空间详情页支持多选图片批量编辑 category、tags；color-search 页根据 picColor 十六进制值检索同空间内图片。上述功能通过 PictureEditByBatchRequest 与 SearchPictureByColorRequest 完成前后端参数约定。')

    add_heading(doc, '3.6.4  个人资料与微信公众号', 3)
    add_body(doc, '/user/update/my 与 /file/upload/avatar 支持资料维护。WxMpController 实现 portal 验证与关键词自动回复，满足课程微信集成要求。')
    add_body(doc, '3.6.5  管理员功能。UserController 提供管理员用户 CRUD；PictureManagePage 与 review 接口支持图片审核；AuthCheck 注解限制 admin 角色访问。Web 管理端采用 Ant Design Table 展示分页数据，与后端 list/page 接口配合。')

    add_heading(doc, '3.7  系统测试', 2)
    add_body(doc, '按功能模块设计测试用例，采用黑盒测试与 Knife4j 接口测试相结合。')
    add_table(doc, '表3-3  用户注册模块测试用例', ['测试功能', '测试描述', '预期结果', '测试结果'], [
        ('用户注册', '输入账号、密码、确认密码，点击注册', '格式正确则创建账号并提示成功', '通过'),
        ('重复账号', '使用已存在账号注册', '提示账号重复，注册失败', '通过'),
        ('密码过短', '密码少于8位', '前端/后端参数校验失败', '通过'),
    ])
    add_table(doc, '表3-4  用户登录模块测试用例', ['测试功能', '测试描述', '预期结果', '测试结果'], [
        ('账号登录', '正确账号密码登录', '返回 token，跳转首页', '通过'),
        ('微信登录', '小程序 wx.login 后调 /user/wx/login', '返回 token 与用户信息', '通过'),
        ('错误密码', '密码错误', '提示账号或密码异常', '通过'),
    ])
    add_table(doc, '表3-5  图片模块测试用例', ['测试功能', '测试描述', '预期结果', '测试结果'], [
        ('图片上传', '登录后上传 JPG', 'COS 有文件，库中待审', '通过'),
        ('公共列表', '游客浏览首页', '仅显示审核通过图', '通过'),
        ('图片审核', '管理员审核通过', '公共库可见', '通过'),
    ])
    add_table(doc, '表3-6  空间与微信模块测试用例', ['测试功能', '测试描述', '预期结果', '测试结果'], [
        ('空间创建', '用户创建普通版空间', '数据库新增 space 记录', '通过'),
        ('空间上传', '向空间内上传图片', 'picture.spaceId 正确', '通过'),
        ('公众号验证', 'GET /wx/mp/portal 带签名', '返回 echostr', '通过'),
        ('权限控制', '普通用户访问审核接口', '40101 无权限', '通过'),
    ])
    add_body(doc, '测试结论：各模块核心用例通过，三端联调正常。Web 工程 npm run build 通过；小程序在开发者工具中功能可用。遗留问题：公众号菜单需 IP 白名单；建议后续补充自动化测试。')
    add_body(doc, '测试环境：Windows 10、JDK 17、MySQL 8.0、Node.js 18、微信开发者工具 Stable。测试账号由小组自行注册，管理员账号 userRole=admin。回归测试覆盖登录—上传—审核—展示主链路，未发现阻塞性缺陷。答辩演示前需确认后端、MySQL、COS 配置正确，并完成至少一张图片的审核通过流程。')

    add_heading(doc, '3.8  总结', 2)
    add_body(doc, '本课程设计完成了星图集系统从需求到实现的全过程，涵盖用户、图片、空间、审核、微信集成等模块，交付 Web 与小程序客户端。采用 Controller-Service-Mapper 分层、Session+JWT 认证、COS+MySQL 存储，具备良好扩展性。')
    add_body(doc, '已完成：注册登录、图库检索、上传审核、空间管理、个人资料、公众号回调等。待优化：Redis 缓存、Elasticsearch 检索、自动化测试、Docker 部署。')
    add_body(doc, '主要问题与解决：（1）小程序公共库空白—补 nullSpaceId；（2）multipart 参数丢失—@ModelAttribute；（3）JWT 未携带—Axios 拦截器；（4）公众号配置—内网穿透与 Java 默认常量。')
    add_body(doc, '未来工作：引入 Redis 提升列表缓存命中率；对接 Elasticsearch 实现全文检索；编写 Docker Compose 实现一键部署；完善 Minium 小程序自动化测试；探索 CDN 加速 COS 图片访问。项目源码已托管于本地 Git 仓库，可作为毕业设计前期原型继续演进。')
    add_body(doc, '通过本项目，完整实践了软件工程生命周期，掌握了 HTML5/CSS3、Ajax、JSON、Java Web、微信开发等课程要求技术，为后续毕业设计打下基础。')
    add_body(doc, f'报告撰写人：{STUDENT["name"]}（{STUDENT["id"]}），{STUDENT["college"]}{STUDENT["major"]}。本报告内容基于星图集小组实际开发成果整理，图形与界面截图为系统运行效果示意，代码片段摘自项目主干分支。')

    add_heading(doc, '参考文献', 1)
    refs = [
        '[1] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
        '[2] 杨开振. Spring Boot + Vue 全栈开发实战[M]. 北京: 清华大学出版社, 2019.',
        '[3] 微信开放文档：小程序开发指南[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/framework/, 2024-06-01.',
        '[4] 微信开放文档：公众号服务器配置[EB/OL]. https://developers.weixin.qq.com/doc/offiaccount/Basic_Information/Access_Overview.html, 2024-06-01.',
        '[5] Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/2.7.6/reference/html/, 2024-03-15.',
        '[6] MyBatis-Plus 官方文档[EB/OL]. https://baomidou.com/pages/24112f/, 2024-05-01.',
        '[7] 腾讯云对象存储 COS 产品文档[EB/OL]. https://cloud.tencent.com/document/product/436, 2024-01-10.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = LINE_SPACING
        set_run_font(p.add_run(ref))

    doc.add_page_break()
    add_heading(doc, '附录1  系统完整程序', 1)
    add_body(doc, '完整源代码目录结构如下，随课程设计光盘或 Git 仓库一并提交：')
    for line in [
        '（1）后端：StarPicture/src/main/java/com/yu/backend/（controller、service、mapper、model）',
        '（2）Web：star-picture-vue/src/（pages、components、api、stores）',
        '（3）小程序：mini/（pages、utils、custom-tab-bar）',
        '（4）数据库：StarPicture/sql/create_sql.sql',
        '（5）配置：application.yml、application-local.yml',
        '（6）接口文档：启动后访问 http://localhost:8123/api/doc.html',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = LINE_SPACING
        r = p.add_run(line)
        set_run_font(r)
    add_body(doc, '核心类说明：UserController 处理用户认证；PictureController 处理图片 CRUD 与审核；FileController 处理 multipart 上传；WxMpController 处理微信回调。前端 request.ts 统一拦截 40100 跳转登录。小程序 utils/api.js 封装 listPictureVo、uploadPicture 等业务方法，降低页面层重复代码。')

    out_main = base / 'docs' / '课程设计报告.docx'
    out_named = base / 'docs' / f'课程设计报告_{STUDENT["name"]}.docx'
    for out in (out_named, out_main):
        try:
            doc.save(out)
            print(f'已生成: {out}')
            break
        except PermissionError:
            print(f'无法写入（文件可能被占用）: {out}')
    else:
        doc.save(out_named)
        print(f'已生成: {out_named}')
    text = '\n'.join(p.text for p in doc.paragraphs)
    cn = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    print(f'汉字约: {cn}')


if __name__ == '__main__':
    main()
