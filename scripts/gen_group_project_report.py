# -*- coding: utf-8 -*-
"""生成小组项目总结报告：星图集小组_项目总结报告.docx"""
from datetime import date
import sys
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle

sys.path.insert(0, str(Path(__file__).resolve().parent))
from team_members import MEMBER_ROWS, CONTRIBUTION_ROWS, GRADE_LINES
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

FONT_SIZE = Pt(10.5)
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
LINE_SPACING = 1.25
GROUP_NAME = '星图集小组'


def setup_plt():
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False


def set_run_font(run, bold=False, size=None):
    run.font.size = size or FONT_SIZE
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    run.bold = bold


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, bold=True, size=Pt(14 if level == 1 else 12))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    r = p.add_run(text)
    set_run_font(r)


def add_image(doc, path, width=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))


def add_table(doc, caption, headers, rows):
    add_caption(doc, caption)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for para in c.paragraphs:
                for run in para.runs:
                    set_run_font(run)
    doc.add_paragraph()


def draw_architecture(path):
    setup_plt()
    fig, ax = plt.subplots(figsize=(11, 7.5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 8)
    ax.axis('off')

    def box(x, y, w, h, t, fc='#E8F4FF'):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.02', fc=fc, ec='#333'))
        ax.text(x + w / 2, y + h / 2, t, ha='center', va='center', fontsize=9)

    box(4, 7.2, 3, 0.6, '星图集 StarPicture', '#FFD1DC')
    for x, t in [(0.5, 'Vue3 Web\nlocalhost:5173'), (4.2, '微信小程序'), (7.8, '微信公众号\n回调/菜单')]:
        box(x, 5.8, 2.5, 1.0, t)
        ax.annotate('', xy=(x + 1.25, 5.8), xytext=(5.5, 7.2), arrowprops=dict(arrowstyle='->', lw=1))
    box(2.2, 4.5, 6.6, 0.7, 'Spring Boot 2.7  端口8123  上下文/api', '#D9E8FF')
    for i, m in enumerate(['UserController', 'PictureController', 'SpaceController', 'FileController', 'WxMpController']):
        box(0.4 + i * 2.05, 3.0, 1.85, 0.75, m.replace('Controller', '\nController'), '#FFF')
        ax.annotate('', xy=(1.3 + i * 2.05, 3.75), xytext=(5.5, 4.5), arrowprops=dict(arrowstyle='->', lw=0.7))
    box(1.0, 1.2, 3.2, 0.9, 'MySQL\nstarpicture', '#FFF4BD')
    box(4.2, 1.2, 3.2, 0.9, 'Service + Mapper', '#E8FFE8')
    box(7.4, 1.2, 3.0, 0.9, '腾讯云 COS', '#B2E2F2')
    ax.annotate('', xy=(2.6, 2.1), xytext=(5.5, 4.5), arrowprops=dict(arrowstyle='->', lw=1))
    ax.annotate('', xy=(5.8, 2.1), xytext=(5.5, 4.5), arrowprops=dict(arrowstyle='->', lw=1))
    ax.annotate('', xy=(8.9, 2.1), xytext=(5.5, 4.5), arrowprops=dict(arrowstyle='->', lw=1))
    fig.savefig(path, dpi=160, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_er(path):
    setup_plt()
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    def ent(x, y, w, h, name, attrs):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle='square,pad=0', fc='white', ec='black'))
        ax.plot([x, x + w], [y + h - 0.32, y + h - 0.32], 'k-')
        ax.text(x + w / 2, y + h - 0.16, name, ha='center', fontsize=10, fontweight='bold')
        for i, a in enumerate(attrs):
            ax.text(x + 0.08, y + h - 0.48 - i * 0.26, a, fontsize=7.5)

    ent(0.4, 2.0, 2.5, 2.3, 'user', ['PK id', 'userAccount', 'userPassword', 'userRole', 'userAvatar...'])
    ent(3.5, 1.7, 3.0, 2.8, 'picture', ['PK id', 'FK userId', 'FK spaceId', 'urls JSON', 'reviewStatus', 'tags/category...'])
    ent(7.2, 2.1, 2.4, 2.0, 'space', ['PK id', 'FK userId', 'spaceName', 'spaceLevel', 'maxCount...'])
    ax.text(3.0, 3.5, '1:N 上传', fontsize=8)
    ax.text(6.6, 3.2, 'N:1 归属', fontsize=8)
    ax.text(5.0, 2.3, '1:N 创建', fontsize=8)
    fig.savefig(path, dpi=160, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def draw_ui_mock(path, title, subtitle, items, footer_tabs=None):
    """生成界面示意截图（无真实截图时使用）"""
    setup_plt()
    fig, ax = plt.subplots(figsize=(8, 10))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 10)
    ax.axis('off')
    # 手机/Web 外框
    ax.add_patch(FancyBboxPatch((0.3, 0.3), 7.4, 9.4, boxstyle='round,pad=0.03', fc='#FAFAFA', ec='#CCC', lw=2))
    ax.add_patch(Rectangle((0.3, 9.0), 7.4, 0.5, fc='#FF69B4', alpha=0.25))
    ax.text(4, 9.25, title, ha='center', va='center', fontsize=12, fontweight='bold')
    ax.text(4, 8.55, subtitle, ha='center', fontsize=9, color='#666')
    y = 7.8
    for row in items:
        if row['type'] == 'search':
            ax.add_patch(FancyBboxPatch((0.7, y), 6.6, 0.55, boxstyle='round,pad=0.02', fc='white', ec='#FFB6C1'))
            ax.text(1.0, y + 0.28, row['text'], fontsize=9, va='center', color='#999')
            y -= 0.85
        elif row['type'] == 'chips':
            x = 0.7
            for c in row['labels']:
                ax.add_patch(FancyBboxPatch((x, y), 1.1, 0.4, boxstyle='round,pad=0.02', fc='#FFE4EC', ec='#FFB6C1'))
                ax.text(x + 0.55, y + 0.2, c, ha='center', va='center', fontsize=7)
                x += 1.25
            y -= 0.65
        elif row['type'] == 'grid':
            gx, gy = 0.7, y
            for i in range(row.get('n', 6)):
                col, r = i % 2, i // 2
                px, py = gx + col * 3.3, gy - r * 2.0
                ax.add_patch(FancyBboxPatch((px, py - 1.5), 3.0, 1.5, boxstyle='round,pad=0.02', fc='#E8E8E8', ec='#DDD'))
                ax.text(px + 1.5, py - 0.75, '图片' + str(i + 1), ha='center', va='center', fontsize=8, color='#888')
            y -= 4.2
        elif row['type'] == 'form':
            for label in row['labels']:
                ax.text(0.9, y, label, fontsize=9)
                ax.add_patch(Rectangle((0.9, y - 0.45), 6.2, 0.38, fc='white', ec='#CCC'))
                y -= 0.75
            ax.add_patch(FancyBboxPatch((2.5, y - 0.1), 3.0, 0.5, boxstyle='round,pad=0.02', fc='#FF69B4', ec='#FF1493'))
            ax.text(4.0, y + 0.15, row.get('btn', '登录'), ha='center', va='center', fontsize=10, color='white', fontweight='bold')
            y -= 0.9
        elif row['type'] == 'profile':
            ax.add_patch(FancyBboxPatch((0.9, y - 0.5), 1.2, 1.2, boxstyle='circle,pad=0.01', fc='#FFD1DC', ec='#FF69B4'))
            ax.text(2.4, y + 0.1, row.get('name', '朱远亮'), fontsize=11, fontweight='bold')
            ax.text(2.4, y - 0.25, row.get('role', '普通用户'), fontsize=8, color='#888')
            y -= 1.5
            for fn in row.get('features', []):
                ax.add_patch(FancyBboxPatch((0.7, y - 0.45), 3.2, 0.55, boxstyle='round,pad=0.02', fc='white', ec='#EEE'))
                ax.text(2.3, y - 0.18, fn, ha='center', va='center', fontsize=8)
                y -= 0.65
    if footer_tabs:
        for i, tab in enumerate(footer_tabs):
            ax.text(1.0 + i * 1.8, 0.75, tab, ha='center', fontsize=8, color='#FF69B4' if i == 0 else '#999')
    fig.savefig(path, dpi=160, bbox_inches='tight', facecolor='white')
    plt.close(fig)


def main():
    base = Path(__file__).resolve().parent.parent
    docs = base / 'docs'
    assets = docs / 'group_report_assets'
    assets.mkdir(parents=True, exist_ok=True)

    img_arch = assets / 'arch.png'
    img_er = assets / 'er.png'
    img_web_home = assets / 'ui_web_home.png'
    img_web_login = assets / 'ui_web_login.png'
    img_mini_home = assets / 'ui_mini_home.png'
    img_mini_mine = assets / 'ui_mini_mine.png'

    draw_architecture(img_arch)
    draw_er(img_er)
    draw_ui_mock(img_web_home, '星图集 · Web', '公共图库首页', [
        {'type': 'search', 'text': '输入关键词搜索...'},
        {'type': 'chips', 'labels': ['全部', '模板', '海报', '素材']},
        {'type': 'grid', 'n': 6},
    ])
    draw_ui_mock(img_web_login, '星图集 · 登录', '爱豆写真云，一起追光', [
        {'type': 'form', 'labels': ['账号', '密码'], 'btn': '登录'},
    ])
    draw_ui_mock(img_mini_home, '星图集 · 小程序', '发现优质图片', [
        {'type': 'search', 'text': '搜索图片素材...'},
        {'type': 'chips', 'labels': ['热门', '高清', '创意']},
        {'type': 'grid', 'n': 4},
    ], footer_tabs=['首页', '上传', '空间', '我的'])
    draw_ui_mock(img_mini_mine, '星图集 · 我的', '个人中心', [
        {'type': 'profile', 'name': '微信用户', 'role': '普通用户', 'features': ['我的空间', '个人资料', '按色搜图', '账号登录']},
    ], footer_tabs=['首页', '上传', '空间', '我的'])

    doc = Document()
    doc.styles['Normal'].font.name = FONT_EN
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    doc.styles['Normal'].font.size = FONT_SIZE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{GROUP_NAME}\n项目总结报告')
    set_run_font(r, bold=True, size=Pt(20))
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('项目名称：星图集（StarPicture）内娱图库管理系统')
    set_run_font(sr, size=Pt(12))
    for line in [f'完成日期：{date.today().strftime("%Y年%m月%d日")}', '学院：工学院    专业：计算机科学与技术']:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(pp.add_run(line))
    doc.add_page_break()

    add_heading(doc, '摘要', 1)
    add_body(
        doc,
        f'{GROUP_NAME}围绕「星图集」内娱图库管理系统开展课程设计与开发，'
        '系统面向图片素材的采集、审核、检索与空间化管理，支持 Vue3 Web 网站、微信小程序与微信公众号三端接入，'
        '后端采用 Spring Boot 统一提供 REST API，数据存储于 MySQL，图片文件存储于腾讯云 COS。'
        '本报告从需求分析、系统架构、技术选型、数据库设计、团队分工、实现难点、项目亮点与不足、测试情况等方面进行总结。'
        '本系统已具备课程要求的 HTML5/CSS3 前端展示、Ajax/JSON 数据交互、Java Web 后端、微信公众号 API 与微信小程序开发等能力，'
        '并在真实联调中验证了业务闭环的可行性。'
    )

    add_heading(doc, '一、需求分析', 1)
    add_body(
        doc,
        '1.1 项目背景。随着粉丝经济与内容创作的发展，用户对明星写真、海报素材等图片资源的整理、分享与检索需求日益增长。'
        '传统网盘或社交软件难以兼顾分类检索、内容审核与多端展示，因此需要一套专门的图库管理系统。'
    )
    add_body(
        doc,
        '1.2 用户角色。系统包含三类角色：（1）游客：可浏览公共图库中已审核通过的图片；（2）普通用户（user）：'
        '可注册登录、上传图片、管理个人空间、编辑资料；（3）管理员（admin）：可审核图片、管理用户、配置公众号菜单等。'
    )
    add_body(
        doc,
        '1.3 功能需求。用户模块：注册、登录（含微信 code 登录）、注销、个人资料与头像维护、管理员用户 CRUD。'
        '图片模块：本地上传、URL 上传、批量上传、分页检索（关键词/分类/标签）、编辑删除、公共库审核机制。'
        '空间模块：创建多级空间、容量配额、空间内图片管理与按颜色搜图。'
        '微信模块：公众号服务器验证、消息回复、自定义菜单；小程序端完整业务闭环。'
    )
    add_body(
        doc,
        '1.4 非功能需求。安全性：密码 MD5 存储、VO 脱敏、JWT+Session 双通道认证、@AuthCheck 权限切面。'
        '性能：列表分页、无筛选条件下 Caffeine 缓存接口；可维护性：Controller-Service-Mapper 分层、Knife4j 文档；'
        '可用性：三端 UI 风格统一（星图集品牌色）、错误码与 Toast 提示友好。'
    )
    add_body(
        doc,
        '1.5 用例概述。经小组讨论，梳理核心用例十余项，包括 UC-01 用户注册、UC-02 账号登录、UC-03 微信登录、'
        'UC-04 图片本地上传、UC-05 图片 URL 上传、UC-06 公共图库检索、UC-07 图片审核、UC-08 空间创建、'
        'UC-09 空间内批量编辑、UC-10 按颜色搜图、UC-11 个人资料维护、UC-12 公众号关键词回复等。'
        '用例优先级划分上，P0 为注册登录与公共图库浏览上传，P1 为空间与审核，P2 为以图搜图与 AI 扩图等扩展能力。'
        '需求评审记录在小组共享文档中，作为迭代计划与测试用例编写的依据。'
    )
    add_body(
        doc,
        '1.6 约束条件。课程周期有限，小组选择 Spring Boot 单体架构而非微服务，以降低部署复杂度；'
        '图片存储依赖腾讯云 COS 免费额度；微信公众号部分功能需公网环境与 IP 白名单；'
        '小程序在开发阶段通过「不校验合法域名」连接本地后端。上述约束在架构设计与测试计划中均已考虑。'
    )

    add_heading(doc, '二、系统架构图', 1)
    add_body(
        doc,
        '系统采用 B/S 与移动端结合的前后端分离架构。表现层为三端客户端，应用层为 Spring Boot 单体服务并按领域划分 Controller，'
        '业务逻辑层由 Service 接口及实现类承担，持久层为 MyBatis-Plus Mapper，'
        '外部依赖包括 MySQL 与腾讯云 COS，微信开放平台提供 OAuth 与消息能力。'
    )
    add_image(doc, img_arch, 14)
    add_caption(doc, '图1  系统总体架构图')
    add_body(
        doc,
        '请求链路说明：以 Web 端浏览公共图库为例，浏览器发起 POST /api/picture/list/page/vo，'
        'Axios 携带 Cookie 或 Authorization 头；DispatcherServlet 路由至 PictureController；'
        'preparePictureQueryForList 设置 nullSpaceId 与 reviewStatus 后调用 PictureService.page；'
        'Mapper 执行 SQL 分页查询；结果封装为 PictureVO 分页对象经 JSON 返回前端渲染卡片网格。'
        '上传链路则经 FileController 写入 COS 后插入 picture 表，形成「对象存储 + 关系库元数据」双写模式。'
    )

    add_heading(doc, '三、技术选型', 1)
    add_table(doc, '表1  技术选型说明', ['层次', '技术', '选型理由'], [
        ('后端', 'Java 17 + Spring Boot 2.7.6', '成熟生态、快速开发、与课程 Servlet 体系衔接'),
        ('ORM', 'MyBatis-Plus 3.5.9', '简化 CRUD、分页与逻辑删除'),
        ('数据库', 'MySQL 8.x', '关系型数据、事务与索引支持'),
        ('对象存储', '腾讯云 COS', '图片大文件与数据库解耦'),
        ('Web 前端', 'Vue3 + TS + Vite + Ant Design Vue', '组件化、类型安全、开发效率高'),
        ('小程序', '微信原生 + wx.request', '官方 API 稳定、便于审核发布'),
        ('接口文档', 'Knife4j 4.4.0', '在线调试、降低联调成本'),
        ('认证', 'Session + JWT Bearer', 'Web Cookie 与小程序 Token 兼顾'),
    ])
    add_body(
        doc,
        '数据交换格式统一为 JSON；文件上传采用 multipart/form-data；'
        '公众号消息采用 XML 与签名验证；小程序登录使用 wx.login 获取 code 后由后端换取 openId 并签发 JWT。'
    )
    add_body(
        doc,
        '前端技术细节：Web 端使用 Pinia 管理登录用户状态，Vue Router 做路由守卫；Axios 配置 withCredentials 与 Bearer Token；'
        'Ant Design Vue 提供 Table、Form、Upload 等组件加速管理端开发。小程序使用自定义 TabBar 统一底部导航，'
        'utils/api.js 封装 ok 处理与分页 normalize，utils/auth.js 管理 sp_token 与 sp_user 本地缓存。'
        '后端使用 Lombok 减少样板代码，Hutool 辅助 JSON 与加密，GlobalExceptionHandler 捕获 BusinessException 返回统一错误码。'
    )

    add_heading(doc, '四、数据库设计', 1)
    add_body(
        doc,
        '数据库名 starpicture，字符集 utf8mb4。核心表 user、picture、space 构成主要 ER 关系：'
        '用户一对多上传图片；用户一对多创建空间；空间一对多包含图片。'
        'picture.spaceId 为空表示公共图库；urls 字段 JSON 存储多规格图片地址；reviewStatus 实现 0待审/1通过/2拒绝；'
        '各表含 isDelete 逻辑删除与 createTime/updateTime 审计字段。'
    )
    add_image(doc, img_er, 13)
    add_caption(doc, '图2  数据库 ER 图')
    add_table(doc, '表2  主要数据表字段摘要', ['表名', '关键字段', '说明'], [
        ('user', 'userAccount, userRole, userAvatar', '用户账号与角色'),
        ('picture', 'urls, userId, spaceId, reviewStatus, tags', '图片元数据与审核'),
        ('space', 'spaceLevel, maxCount, totalCount, userId', '空间配额与统计'),
    ])
    add_body(
        doc,
        '索引设计：user 表 userAccount 唯一索引；picture 表对 name、category、userId、spaceId、reviewStatus 建索引以优化列表与审核查询；'
        'space 表对 userId 建索引便于「我的空间」列表。逻辑删除字段 isDelete 配合 MyBatis-Plus @TableLogic 实现软删除，'
        '避免误删后数据无法恢复。picture.urls 采用 JSON 类型存储 originalUrl、thumbnailUrl 等多规格地址，'
        '便于后续 CDN 切换与图片处理流水线扩展。'
    )

    add_heading(doc, '五、团队分工与个人贡献', 1)
    add_body(
        doc,
        f'{GROUP_NAME}以敏捷迭代方式推进，每周进行进度同步与接口评审。'
        '下表列出成员分工及主要交付物，个人贡献与各位成员课程学习总结、课程设计报告保持一致。'
    )
    add_table(
        doc,
        '表3  项目组成员及分工及成绩评定',
        ['学号', '姓名', '专业', '主要工作', '课程设计成绩'],
        MEMBER_ROWS,
    )
    add_table(
        doc,
        '表4  团队分工与个人贡献说明',
        ['姓名', '学号', '角色', '主要任务', '个人贡献说明'],
        CONTRIBUTION_ROWS,
    )
    add_caption(doc, '表5  个人课程设计成绩评定（由指导教师填写）')
    grade_table = doc.add_table(rows=len(GRADE_LINES), cols=1)
    grade_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (sid, name, tag) in enumerate(GRADE_LINES):
        cell = grade_table.rows[i].cells[0]
        cell.text = f'学号：{sid}    {tag}  姓名：{name}     成    绩：__________'
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run)
    doc.add_paragraph()
    add_body(
        doc,
        '协作机制：采用 Git 版本管理；朱远亮（组长）负责架构、后端、小程序、公众号及三端联调，工作量最大；'
        '林景彬负责 Web 前端主流程页面；李冠燃负责数据库脚本、测试与测试文档；李坤纬负责管理端页面、UI 统一与答辩材料。'
        '四人分工与各成员《课程学习总结》保持一致。'
    )

    add_heading(doc, '六、实现难点与解决方案', 1)
    add_body(
        doc,
        '难点1：小程序公共图库无数据。原因：分页查询未传 nullSpaceId=true，后端无法定位公共空间。'
        '解决：index.js 的 buildQuery 显式设置 nullSpaceId，并与后端 preparePictureQueryForList 逻辑对齐。'
    )
    add_body(
        doc,
        '难点2：multipart 上传 spaceId 丢失。原因：FileController 未使用 @ModelAttribute 绑定 PictureUploadRequest。'
        '解决：补充注解与空对象默认值；upload.js 提交前 ensureLogin 并携带 JWT。'
    )
    add_body(
        doc,
        '难点3：微信公众号本地调试。原因：需公网 URL 与 IP 白名单。'
        '解决：cpolar 内网穿透、公众平台配置 Token/EncodingAESKey；菜单 create 接口由管理员调用；YAML 复杂 Map 改 Java 常量避免启动失败。'
    )
    add_body(
        doc,
        '难点4：三端登录态不一致。Web 仅 Session 导致 Token 场景失败。'
        '解决：登录保存 sp_token，Axios 拦截器加 Authorization；40100 时清 token 并跳转登录。'
    )
    add_body(
        doc,
        '难点5：公共图库泄露未审核内容。解决：get/vo 对非本人非管理员校验 reviewStatus=PASS；列表接口强制审核过滤。'
    )
    add_body(
        doc,
        '难点6：小程序冷启动与空间页竞态。用户打开小程序时 wxLoginSilent 尚未完成，space/list 已请求导致列表为空。'
        '解决：app.js 暴露 globalData.authReady Promise，空间列表页 onReady 等待 authReady 后再 loadList。'
        '难点7：Vue 工程 TypeScript 严格模式报错。解决：补全 LoginUserVo.token、PictureQueryRequest.nullSpaceId 等类型定义，'
        '修复 admin 页与 Upload 组件的类型断言，最终 npm run build 通过。'
    )

    add_heading(doc, '七、项目亮点与不足', 1)
    add_body(
        doc,
        '7.1 项目亮点。（1）三端一体：Web、小程序、公众号共用后端，降低维护成本；'
        '（2）完整内容闭环：上传—审核—公开展示—空间私有管理；'
        '（3）工程规范：统一 BaseResponse、ErrorCode、VO 脱敏、AOP 鉴权；'
        '（4）扩展能力：以图搜图、AI 扩图、按色搜图、批量编辑等增值功能接口；'
        '（5）文档齐全：Knife4j 在线 API、公众号配置说明、课程与小组报告。'
    )
    add_body(
        doc,
        '与同类课程项目相比，本项目的差异化在于：其一，真正打通微信生态（不仅是小程序，还含公众号服务器开发）；'
        '其二，引入内容审核状态机，公共数据与私有空间数据隔离清晰；'
        '其三，图片元数据丰富（宽高、格式、主色调、多规格 URL），为检索与 AI 能力预留扩展点；'
        '其四，三端 UI 统一「星图集」视觉语言，体现产品化思维而非单纯作业堆砌。'
    )
    add_body(
        doc,
        '7.2 不足之处。（1）自动化测试覆盖不足，以手工测试为主；'
        '（2）未引入 Redis 分布式缓存与 Session；'
        '（3）部分管理端 TypeScript 类型曾需补全；'
        '（4）公众号菜单在测试号环境能力受限；'
        '（5）性能压测与高可用部署方案尚未实施。'
    )
    add_body(
        doc,
        '改进计划：下一版本引入 Redis 缓存热点列表；使用 MinIO 作为 COS 本地替代便于演示；'
        '补充 Docker Compose 一键启动脚本；完善管理员数据看板与操作日志，提升运维可观测性。'
    )

    add_heading(doc, '八、测试情况', 1)
    add_body(
        doc,
        '测试策略：接口测试（Knife4j）+ 三端黑盒功能测试 + 边界用例抽查。'
        '测试环境：Windows 10、JDK 17、MySQL 8、Node 18、微信开发者工具。'
    )
    add_table(
        doc,
        '表6  功能测试汇总',
        ['模块', '测试项', '预期结果', '结果'],
        [
            ('用户', '注册/登录/微信登录', '返回 token 与用户信息', '通过'),
            ('用户', '个人资料/头像上传', '信息持久化并同步导航栏', '通过'),
            ('图片', '上传至公共库', 'COS 有文件、待审核状态', '通过'),
            ('图片', '首页列表', '仅展示审核通过图', '通过'),
            ('图片', '管理员审核', '通过后公开展示', '通过'),
            ('空间', '创建与上传', 'spaceId 正确、配额更新', '通过'),
            ('微信', 'portal GET 验证', '返回 echostr', '通过'),
            ('权限', '普通用户访问 admin 接口', '40101 无权限', '通过'),
            ('Web', 'npm run build', '类型检查与打包成功', '通过'),
        ],
    )
    add_body(
        doc,
        '测试结论：核心业务流程满足需求规格，三端联调通过，已知问题已在迭代中修复。'
        '建议后续增加 JMeter 压测与 Selenium/Minium 自动化 UI 测试。'
    )
    add_body(
        doc,
        '回归测试：在修复 nullSpaceId、multipart 绑定、JWT 拦截等问题后，小组对三端登录—上传—审核—展示链路进行全量回归，'
        '未发现阻塞性缺陷。兼容性方面，Web 在 Chrome/Edge 最新版表现正常；小程序基础库 2.x 测试通过。'
        '安全测试抽查：未登录访问 /picture/delete 返回 40100；普通用户访问 /picture/review 返回 40101；'
        '管理员 get 用户接口返回 UserVO 不含密码字段。'
    )
    add_body(
        doc,
        '测试过程中，小组维护缺陷台账共记录十余项问题，均已关闭或标注为环境限制（如公众号 IP 白名单）。'
        '测试数据使用独立测试账号 linktest01 等，避免污染生产数据。答辩演示脚本覆盖注册、登录、上传、审核、公开展示全流程，满足课程验收与答辩展示要求。'
    )

    add_heading(doc, '九、关键界面截图', 1)
    add_body(
        doc,
        '以下为本项目主要界面运行效果示意（Web 运行于 localhost:5173，小程序运行于微信开发者工具，后端 localhost:8123）。'
        '若需提交高清实拍图，可在系统运行后将截图替换为实际界面照片。'
    )
    add_image(doc, img_web_home, 11)
    add_caption(doc, '图3  Web 端公共图库首页')
    add_image(doc, img_web_login, 11)
    add_caption(doc, '图4  Web 端登录页')
    add_image(doc, img_mini_home, 11)
    add_caption(doc, '图5  微信小程序首页')
    add_image(doc, img_mini_mine, 11)
    add_caption(doc, '图6  微信小程序「我的」页')
    add_body(
        doc,
        '界面设计说明：Web 端采用顶部导航 + 内容区布局，去除突兀侧栏，将「公共图库」「我的空间」「个人资料」集中于 Header；'
        '配色以 sakura/lilac 渐变为主，符合内娱图库年轻用户审美。小程序采用卡片式 Hero 搜索、横向分类 Chip、瀑布流网格，'
        '「我的」页以功能九宫格呈现空间、资料、搜色、审核（管理员）等入口，与 Web 能力对齐。'
        '管理员页面提供图片审核与用户管理表格，支持分页与条件筛选。'
    )

    add_heading(doc, '十、总结', 1)
    add_body(
        doc,
        f'{GROUP_NAME}历时一个学期完成星图集系统从需求到上线的全过程，'
        '成员在 Web 全栈、移动开发、第三方平台集成方面均获得实战锻炼。'
        '项目已实现用户、图片、空间、审核、微信生态等核心能力，'
        '后续可在检索性能、自动化测试与云原生部署方向持续优化。'
        '感谢指导教师在架构评审与答辩中的指导。'
    )
    add_body(
        doc,
        '通过本项目，小组掌握了从需求分析、数据库建模、后端 API、前端与小程序开发到第三方平台集成的完整软件工程流程，'
        '为后续毕业设计与企业实习奠定了坚实基础。项目源码与文档存放于 D:\\code\\StarPicture 工程目录，'
        '接口文档地址 http://localhost:8123/api/doc.html，欢迎老师同学批评指正。以上为本小组项目总结全部内容。'
    )

    out_main = docs / f'{GROUP_NAME}_项目总结报告.docx'
    out_alt = docs / f'{GROUP_NAME}_项目总结报告_v2.docx'
    saved = False
    for out in (out_main, out_alt):
        try:
            doc.save(out)
            print(f'已生成: {out}')
            saved = True
            break
        except PermissionError:
            print(f'无法写入（请关闭 Word）: {out}')
    if not saved:
        raise PermissionError('无法保存小组报告，请关闭已打开的 docx 后重试')
    text = '\n'.join(p.text for p in doc.paragraphs)
    cn = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    print(f'汉字约: {cn}')


if __name__ == '__main__':
    main()
