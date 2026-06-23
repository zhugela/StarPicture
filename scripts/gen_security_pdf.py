"""
生成 4 份 PDF 安全报告模板
放到各成员文件夹的 安全测试/ 目录下
"""
from docx import Document
from pathlib import Path
from datetime import datetime

BASE = Path("D:/code/StarPicture/docs/test")
today = "2026-06-19"

def make_report(module, owner, scope, test_items):
    doc = Document()
    title = doc.add_heading(f'内娱图库（StarPicture）安全测试报告', 0)
    title.alignment = 1
    sub = doc.add_heading(f'模块：{module}', 1)
    sub.alignment = 1

    # 元数据
    p = doc.add_paragraph()
    p.add_run('测试人员：').bold = True
    p.add_run(f'{owner}\n')
    p.add_run('测试日期：').bold = True
    p.add_run(f'{today}\n')
    p.add_run('测试范围：').bold = True
    p.add_run(scope)

    # 1. 测试目的
    doc.add_heading('1. 测试目的', 1)
    doc.add_paragraph(
        f'本报告针对内娱图库（StarPicture）的【{module}】进行安全测试，'
        f'验证系统在常见安全攻击下的防御能力，确保系统不存在高危漏洞，'
        f'符合上线安全要求。'
    )

    # 2. 测试环境
    doc.add_heading('2. 测试环境', 1)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Light Grid Accent 1'
    t.rows[0].cells[0].text = '项目'
    t.rows[0].cells[1].text = '内容'
    for k, v in [
        ('目标系统', '内娱图库（StarPicture） V2.00'),
        ('部署环境', '本地开发环境（Windows 11）'),
        ('后端地址', 'http://localhost:8123/api'),
        ('JDK', 'OpenJDK 17.0.10'),
        ('数据库', 'MySQL 8.0.x（含分库分表）'),
        ('测试工具', 'AppScan / BurpSuite / Postman / 手工测试'),
        ('测试账号', 'admin / testuser01 / testuser02 / testuser03'),
    ]:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v

    # 3. 测试项
    doc.add_heading('3. 测试项', 1)
    doc.add_paragraph('本次安全测试覆盖以下测试项：')
    for ti in test_items:
        doc.add_paragraph(ti, style='List Bullet')

    # 4. 风险等级说明
    doc.add_heading('4. 风险等级说明', 1)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    t.rows[0].cells[0].text = '等级'
    t.rows[0].cells[1].text = '说明'
    t.rows[0].cells[2].text = '数量'
    for k, v, c in [
        ('高危', '系统崩溃、数据泄露、绕过认证', '[待填]'),
        ('中危', '重要功能被绕过、有 workaround', '[待填]'),
        ('低危', '轻微信息泄露、UX 问题', '[待填]'),
        ('建议', '最佳实践偏离、可优化', '[待填]'),
    ]:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
        row[2].text = c

    # 5. 测试结果详情
    doc.add_heading('5. 测试结果详情', 1)
    doc.add_paragraph(
        '下表列出本次安全测试的具体结果。'
        '每条测试项都需在测试执行完毕后由测试人员填写实际结论。'
    )
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Light Grid Accent 1'
    t.rows[0].cells[0].text = '编号'
    t.rows[0].cells[1].text = '测试项'
    t.rows[0].cells[2].text = '测试方法'
    t.rows[0].cells[3].text = '风险等级'
    t.rows[0].cells[4].text = '结论'
    for i in range(1, 6):
        row = t.add_row().cells
        row[0].text = f'SEC-{i:03d}'
        row[1].text = '[待填：测试项名称]'
        row[2].text = '[待填：手工/工具扫描/注入测试]'
        row[3].text = '[待填：高/中/低/无]'
        row[4].text = '[待填：通过/不通过，附说明]'

    # 6. 漏洞详情
    doc.add_heading('6. 漏洞详情', 1)
    doc.add_paragraph(
        '若测试中发现漏洞，请按以下格式在每个漏洞前添加一个小标题（Heading 2）并填写详情。'
    )
    doc.add_heading('示例：SEC-001 SQL 注入漏洞', 2)
    doc.add_paragraph('【发现时间】：[待填]')
    doc.add_paragraph('【风险等级】：[待填]')
    doc.add_paragraph('【影响范围】：[待填]')
    doc.add_paragraph('【复现步骤】：[待填]')
    doc.add_paragraph('【修复建议】：[待填]')
    doc.add_paragraph('【当前状态】：[待填：已修复/未修复/延期]')

    # 7. 测试结论
    doc.add_heading('7. 测试结论', 1)
    doc.add_paragraph(
        '【整体评价】：[待填：通过/有条件通过/不通过]\n'
        '【统计】：高危 X 个，中危 X 个，低危 X 个，建议 X 个\n'
        '【是否影响上线】：[待填：是/否]\n'
        '【遗留风险】：[待填：若通过，简述遗留风险；若不通过，说明原因]'
    )

    # 8. 附件
    doc.add_heading('8. 附件', 1)
    doc.add_paragraph('附件 1：AppScan 扫描报告（.scan 文件）')
    doc.add_paragraph('附件 2：Postman 复现脚本（.postman_collection.json）')
    doc.add_paragraph('附件 3：相关截图（*.png）')

    return doc

# ============ 4 份 PDF ============
reports = [
    ("用户模块", "朱远亮", "user 模块（注册/登录/CRUD/认证）", [
        "TC-SEC-001：SQL 注入测试（登录用户名/密码）",
        "TC-SEC-002：越权访问测试（普通用户调用管理员接口）",
        "TC-SEC-003：Cookie 伪造测试",
    ]),
    ("图片模块", "李冠燃", "picture 模块（上传/编辑/查询/搜索）", [
        "TC-SEC-001：木马上传测试（伪 JSP/PHP 文件）",
        "TC-SEC-002：SSRF 测试（URL 上传 + 编辑代理）",
        "TC-SEC-003：XSS 测试（图片 name/introduction/tags）",
    ]),
    ("空间模块", "李坤纬", "space 模块（空间+成员+分析）", [
        "TC-SEC-001：空间名称 XSS 注入",
        "TC-SEC-002：越权添加成员",
    ]),
    ("文件_公众号模块", "林景彬", "file + wxMp 模块（文件上传+微信公众号）", [
        "TC-SEC-001：木马上传（伪 PHP 内容 .jpg 文件）",
        "TC-SEC-002：双扩展名绕过（1.jpg.php）",
        "TC-SEC-003：越权修改他人头像",
    ]),
]

for module, owner, scope, items in reports:
    if module == "用户模块":
        d = BASE / "朱远亮_脚本与截图/安全测试"
    elif module == "图片模块":
        d = BASE / "李冠燃_脚本与截图/安全测试"
    elif module == "空间模块":
        d = BASE / "李坤纬_脚本与截图/安全测试"
    else:
        d = BASE / "林景彬_脚本与截图/安全测试"
    d.mkdir(parents=True, exist_ok=True)

    doc = make_report(module, owner, scope, items)
    # 存为 .docx（用户用 WPS 打开后另存为 PDF）
    docx_path = d / f"{module}-安全报告模板.docx"
    doc.save(docx_path)
    print(f"已生成: {docx_path}")

print(f"\n共 {len(reports)} 份 PDF 报告模板。")
print("说明：模板为 .docx 格式，4 人用 WPS 打开填实际结论后，'文件 → 导出 → PDF' 即可。")
