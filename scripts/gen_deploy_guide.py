# -*- coding: utf-8 -*-
"""生成「星图集 StarPicture 换机部署与全栈配置指南」Word 文档"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_shading(cell, fill: str):
    try:
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    except Exception:
        pass


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], 'D9E8FF')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_code(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def main():
    out = Path(__file__).resolve().parent.parent / 'docs' / '星图集_换机部署与全栈配置指南.docx'
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # ===== 封面 =====
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('星图集 StarPicture\n换机部署与全栈配置指南')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f'文档版本：V1.0    编写日期：{date.today().isoformat()}\n'
        '适用范围：后端 + Vue 网页 + 微信小程序 + 微信公众号'
    )
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        '本文档说明如何在「新电脑」上从零配置并跑通星图集项目的全部模块。'
        '项目采用「一个 Spring Boot 后端 + 三个前端入口」架构：Vue 网页、微信小程序、微信公众号均调用同一套 API 与 MySQL 数据库。'
        '按本文顺序操作，可在本地完成课程演示与联调；正式上线需额外配置 HTTPS 域名与微信服务器白名单。'
    )
    intro.paragraph_format.first_line_indent = Cm(0.74)

    # ===== 1 项目结构 =====
    doc.add_heading('一、项目结构与代码仓库', level=1)
    add_table(doc, ['目录/仓库', '技术栈', '说明'], [
        ['StarPicture', 'Java 17 + Spring Boot 2.7.6 + MyBatis-Plus', '后端 API，端口 8123，上下文 /api'],
        ['star-picture-vue', 'Vue 3 + Vite + Ant Design Vue', 'Web 管理端与公共图库，开发端口 5173'],
        ['mini', '微信小程序原生', '小程序端，微信开发者工具打开'],
    ], [4.5, 5.5, 5.5])

    doc.add_paragraph('建议在新电脑上保持相同目录结构，例如：')
    add_code(doc, '''D:\\code\\StarPicture          # 后端
D:\\code\\star-picture-vue    # 网页前端
D:\\code\\mini                 # 小程序''')

    doc.add_paragraph('迁移方式（任选其一）：')
    add_bullets(doc, [
        'Git 克隆：三台仓库分别 git clone 到新电脑',
        'U 盘/网盘：复制整个 code 文件夹（不要复制 node_modules，到新电脑后重新 npm install）',
        '旧电脑打包：StarPicture 可不含 target；mini 与 vue 不含 node_modules',
    ])

    # ===== 2 环境 =====
    doc.add_heading('二、开发环境安装清单', level=1)
    add_table(doc, ['软件', '推荐版本', '用途', '验证命令'], [
        ['JDK', '17（IDEA 可设 17；pom 兼容 8+）', '运行 Spring Boot', 'java -version'],
        ['Maven', '3.8+', '构建后端', 'mvn -version'],
        ['IntelliJ IDEA', '2023+', '后端开发/启动', '—'],
        ['MySQL', '8.x', '业务数据库 starpicture', 'mysql --version'],
        ['Node.js', '18 LTS 或 20 LTS', 'Vue 前端', 'node -v && npm -v'],
        ['微信开发者工具', '最新稳定版', '小程序编译预览', '—'],
        ['Git', '2.x', '拉取代码', 'git --version'],
        ['（可选）cpolar / ngrok', '—', '公众号本地 HTTPS 穿透', '—'],
    ], [3, 3.5, 5, 4])

    doc.add_paragraph('网络与账号准备：')
    add_bullets(doc, [
        '腾讯云 COS 对象存储账号（图片文件存储，SecretId/SecretKey/桶名/地域）',
        '微信公众平台小程序 AppID + AppSecret（小程序登录）',
        '微信公众平台公众号 AppID + AppSecret（自动回复、自定义菜单；可与小程序不同主体，菜单跳转需关联小程序）',
        '（可选）阿里云 AI API Key：扩图等 AI 功能',
    ])

    # ===== 3 换机总流程 =====
    doc.add_heading('三、换机部署总流程（推荐顺序）', level=1)
    add_numbers(doc, [
        '安装 JDK、Maven、MySQL、Node.js、微信开发者工具',
        '复制/克隆三个项目到本地',
        '创建 MySQL 数据库并导入 sql/starpicture_full.sql（完整一键脚本）',
        '配置 StarPicture/src/main/resources/application-local.yml（数据库、COS、微信密钥）',
        '启动后端，访问 http://127.0.0.1:8123/api/doc.html 确认接口文档可用',
        '在 star-picture-vue 执行 npm install && npm run dev，访问 http://localhost:5173',
        '用微信开发者工具打开 mini 目录，确认 utils/config.js 中 BASE_URL 指向后端',
        '（公众号）配置内网穿透 + 公众平台服务器 URL + 创建自定义菜单',
        '按第七章验收清单逐项测试',
    ])

    # ===== 4 后端 =====
    doc.add_heading('四、后端 StarPicture 配置', level=1)

    doc.add_heading('4.1 启动方式', level=2)
    add_bullets(doc, [
        'IDEA：打开 StarPicture → 运行 Backend2Application.java',
        '命令行：cd StarPicture && mvn spring-boot:run',
        '打包：mvn clean package -DskipTests，java -jar target/backend2-0.0.1-SNAPSHOT.jar',
    ])

    doc.add_heading('4.2 核心配置文件', level=2)
    add_table(doc, ['文件', '说明'], [
        ['src/main/resources/application.yml', '端口 8123、context-path /api、knife4j、wx 框架配置'],
        ['src/main/resources/application-local.yml', '本地敏感配置：MySQL、COS、小程序/公众号密钥（勿提交 Git）'],
    ], [6, 9.5])

    doc.add_paragraph('application-local.yml 必须修改的项（从旧电脑复制或重新填写，勿泄露到公开仓库）：')
    add_code(doc, '''spring:
  datasource:
    url: jdbc:mysql://localhost:3306/starpicture?useSSL=false&serverTimezone=UTC&characterEncoding=utf-8&useUnicode=true&allowPublicKeyRetrieval=true
    username: root          # 改成新电脑 MySQL 用户名
    password: <你的密码>   # 改成新电脑 MySQL 密码

cos:
  client:
    host: https://<桶名>.cos.<地域>.myqcloud.com
    secretId: <腾讯云 SecretId>
    secretKey: <腾讯云 SecretKey>
    region: ap-guangzhou    # 与桶一致
    bucket: <桶名-appid>
    simple-upload-only: true

wx:
  miniapp:
    app-id: wxcc006e978973a40b      # 小程序 AppID（以公众平台为准）
    app-secret: <小程序 AppSecret>
    token-secret: <JWT 密钥，任意长字符串>
  mp:
    enabled: true
    app-id: <公众号 AppID>
    app-secret: <公众号 AppSecret>
    token: starpicture              # 与公众平台服务器 Token 一致''')

    doc.add_heading('4.3 MySQL 数据库', level=2)
    add_numbers(doc, [
        '安装 MySQL 8，记住 root 密码',
        '执行 sql/starpicture_full.sql（推荐，含最新表结构 + 可选演示账号）',
        '若从旧电脑迁数据：Navicat 导出 starpicture 整库 → 新电脑导入（可跳过 full.sql）',
        '历史增量脚本：sql/create_sql.sql、sql/add_user_mp_open_id.sql（已由 full.sql 合并）',
    ])

    doc.add_heading('4.4 腾讯云 COS', level=2)
    doc.add_paragraph(
        '图片上传、URL 导入、缩略图均依赖 COS。在新电脑只需在 application-local.yml 填写与旧环境相同的 '
        'secretId、secretKey、region、bucket、host。若出现 SignatureDoesNotMatch，检查密钥是否与当前腾讯云账号匹配，'
        'region 与 bucket 是否与控制台一致。'
    )

    doc.add_heading('4.5 后端验证', level=2)
    add_table(doc, ['检查项', '地址/方法', '预期结果'], [
        ['接口文档', 'http://127.0.0.1:8123/api/doc.html', 'Knife4j 页面正常打开'],
        ['分类标签', 'GET /api/picture/tag_category', '返回 categoryList、tagList'],
        ['图库列表', 'POST /api/picture/list/page/vo', 'body: {"current":1,"pageSize":12} 返回 records'],
        ['用户注册/登录', 'POST /api/user/register、/login', 'code=0'],
    ], [3.5, 5.5, 6.5])

    # ===== 5 Vue =====
    doc.add_heading('五、Vue 网页前端配置（star-picture-vue）', level=1)

    doc.add_heading('5.1 安装与启动', level=2)
    add_code(doc, '''cd D:\\code\\star-picture-vue
npm install
npm run dev''')

    doc.add_paragraph('浏览器访问：http://localhost:5173')

    doc.add_heading('5.2 接口地址配置', level=2)
    doc.add_paragraph('开发环境通过 Vite 代理转发到本地后端，一般无需改代码：')
    add_code(doc, '''// vite.config.ts
server: {
  proxy: {
    '/api': { target: 'http://localhost:8123', changeOrigin: true },
  },
}''')

    doc.add_paragraph('src/request.ts 中 baseURL 为 http://localhost:8123，与后端端口一致。'
                      '若后端改端口，需同步修改 vite.config.ts 与 request.ts。')

    doc.add_heading('5.3 登录 Token', level=2)
    doc.add_paragraph(
        '网页登录成功后 Token 保存在 localStorage（键名 sp_token），请求头自动带 Authorization: Bearer <token>，'
        '与小程序共用同一套 JWT 鉴权。'
    )

    doc.add_heading('5.4 生产部署（可选）', level=2)
    add_bullets(doc, [
        'npm run build 生成 dist/',
        '将 dist 部署到 Nginx，反向代理 /api 到后端 HTTPS 地址',
        '正式环境须使用 HTTPS，否则微信相关能力受限',
    ])

    # ===== 6 小程序 =====
    doc.add_heading('六、微信小程序配置（mini）', level=1)

    doc.add_heading('6.1 打开项目', level=2)
    add_numbers(doc, [
        '安装并打开「微信开发者工具」',
        '导入项目 → 目录选择 mini 文件夹',
        'AppID 填写 wxcc006e978973a40b（或使用测试号，但 wx.login 需真实 AppSecret）',
        'project.config.json 中 setting.urlCheck 建议开发时设为 false',
    ])

    doc.add_heading('6.2 后端地址 config.js', level=2)
    doc.add_paragraph('文件路径：mini/utils/config.js')
    add_code(doc, '''const USE_LOCALHOST = true   // 开发者工具模拟器：true
const PC_IP = '192.168.1.100' // 真机预览：改为新电脑局域网 IP
const USE_PROD = false
const PROD_BASE_URL = 'https://你的域名.com/api'

// 开发者工具：http://127.0.0.1:8123/api
// 手机预览：http://<PC_IP>:8123/api（手机与电脑同一 WiFi）''')

    add_table(doc, ['场景', 'USE_LOCALHOST', 'PC_IP / PROD', '说明'], [
        ['开发者工具模拟器', 'true', '—', '127.0.0.1:8123，后端须在本机运行'],
        ['手机真机预览', 'false', '填新电脑局域网 IP', 'Windows: ipconfig 查看 IPv4'],
        ['体验版/正式版', '—', 'USE_PROD=true + HTTPS 域名', '须在微信后台配置 request 合法域名'],
    ], [4, 3, 4.5, 4])

    doc.add_heading('6.3 小程序登录', level=2)
    doc.add_paragraph(
        '小程序启动时 app.js 会调用 wx.login → POST /api/user/wx/login。'
        '后端 application-local.yml 必须配置正确的 wx.miniapp.app-id 与 app-secret，'
        '否则登录失败、部分功能不可用。'
    )

    doc.add_heading('6.4 小程序验证', level=2)
    add_bullets(doc, [
        '图库页（pages/index/index）能加载分类标签和图片列表',
        '上传页可选择图片并成功上传（依赖 COS 配置）',
        '我的页可看到登录用户信息',
        '若列表为空：确认后端有已过审且 spaceId 为空的公共图片；首页已改为走非缓存 list 接口',
    ])

    # ===== 7 公众号 =====
    doc.add_heading('七、微信公众号配置', level=1)
    doc.add_paragraph(
        '公众号与小程序共用 StarPicture 后端。详细菜单、自动回复、穿透步骤另见 docs/公众号配置说明.docx；'
        '以下为换机时必须重做/核对的要点。'
    )

    doc.add_heading('7.1 后端配置', level=2)
    add_bullets(doc, [
        'application-local.yml → wx.mp.enabled=true',
        'wx.mp.app-id / app-secret 为公众号凭证（不是小程序 AppID）',
        'wx.mp.token 与公众平台「服务器配置 → Token」完全一致（默认 starpicture）',
    ])

    doc.add_heading('7.2 公众平台服务器 URL', level=2)
    add_table(doc, ['字段', '填写内容'], [
        ['URL', 'https://<公网HTTPS域名>/api/wx/mp/portal'],
        ['Token', 'starpicture（与 yml 一致）'],
        ['消息加解密', '明文模式（当前后端实现）'],
    ], [4, 11.5])

    doc.add_paragraph('本地开发无公网 IP 时，使用 cpolar/ngrok 将 8123 映射为 HTTPS，例如：')
    add_code(doc, 'cpolar http 8123\n# 得到 https://xxxx.cpolar.top\n# 完整 URL：https://xxxx.cpolar.top/api/wx/mp/portal')

    doc.add_heading('7.3 创建自定义菜单', level=2)
    add_numbers(doc, [
        '公众平台「基本配置」添加本机公网 IP 到 IP 白名单',
        'MySQL 将测试账号 userRole 设为 admin',
        'POST /api/user/login 获取 JWT',
        'POST /api/wx/mp/menu/create，Header: Authorization: Bearer <token>',
        '或在 Knife4j → 微信公众号 → 创建默认自定义菜单',
        '公众号后台关联小程序 AppID（wxcc006e978973a40b）',
    ])

    doc.add_heading('7.4 公众号验收', level=2)
    add_bullets(doc, [
        '关注公众号收到欢迎语',
        '发送「帮助」「图库」有自动回复',
        '底部菜单「打开图库」可跳转小程序',
    ])

    # ===== 8 启动顺序 =====
    doc.add_heading('八、日常启动顺序', level=1)
    add_table(doc, ['顺序', '模块', '操作', '访问地址'], [
        ['1', 'MySQL', '确保 MySQL 服务已启动', 'localhost:3306'],
        ['2', '后端', 'IDEA 运行 Backend2Application', 'http://127.0.0.1:8123/api/doc.html'],
        ['3', 'Vue 网页', 'npm run dev', 'http://localhost:5173'],
        ['4', '小程序', '微信开发者工具 → 编译', '模拟器 / 真机预览'],
        ['5', '公众号', '（可选）启动 cpolar + 后端', '手机微信关注测试'],
    ], [1.5, 2.5, 5.5, 5])

    # ===== 9 验收 =====
    doc.add_heading('九、全链路验收清单', level=1)
    checklist = [
        '[ ] MySQL starpicture 库表齐全，能查到 user/picture 数据',
        '[ ] 后端 doc.html 可打开，list/page/vo 有图片记录',
        '[ ] Vue 首页公共图库显示图片，可登录/注册',
        '[ ] Vue 可上传图片、管理空间（管理员可审核）',
        '[ ] 小程序图库页显示图片（非「暂无图片」）',
        '[ ] 小程序可上传、进入详情页',
        '[ ] 小程序 wx.login 登录成功（需 app-secret 正确）',
        '[ ] 公众号服务器 URL 验证通过（若做公众号演示）',
        '[ ] 公众号自动回复与菜单正常（若做公众号演示）',
    ]
    for c in checklist:
        doc.add_paragraph(c)

    # ===== 10 故障 =====
    doc.add_heading('十、常见问题排查', level=1)
    add_table(doc, ['现象', '可能原因', '处理'], [
        ['后端启动失败 / 连不上数据库', 'MySQL 未启动或密码错误', '检查 application-local.yml 的 url/username/password'],
        ['上传图片 COS 403', 'SecretId/Key 或 region/bucket 错误', '对照腾讯云控制台 COS 配置'],
        ['网页有图、小程序无图', '缓存或 BASE_URL 错误', 'config.js 指向 8123；重启后端；小程序重新编译'],
        ['小程序 request 失败', '真机用了 127.0.0.1', '改 PC_IP 为电脑局域网 IP，USE_LOCALHOST=false'],
        ['wx.login 失败', 'app-secret 未配置或 AppID 不匹配', '核对 application-local.yml wx.miniapp'],
        ['公众号 Token 验证失败', 'Token 不一致或 URL 缺 /api', '统一 starpicture；URL 含 /api/wx/mp/portal'],
        ['URL 上传报 Windows 路径错误', '旧版把 URL 当文件名', '使用已修复的后端 UrlPictureUpload 代码'],
        ['npm install 慢/失败', '网络或 Node 版本', '换 Node 18 LTS；可用 npm config set registry 国内镜像'],
    ], [3.5, 4.5, 7.5])

    # ===== 11 附录 =====
    doc.add_heading('十一、关键文件速查表', level=1)
    add_table(doc, ['模块', '配置文件', '作用'], [
        ['后端', 'application-local.yml', '数据库、COS、微信密钥'],
        ['后端', 'application.yml', '端口、context-path、knife4j'],
        ['后端', 'sql/starpicture_full.sql', '新电脑完整建库建表'],
        ['Vue', 'vite.config.ts', '开发代理 /api → 8123'],
        ['Vue', 'src/request.ts', 'axios baseURL'],
        ['小程序', 'utils/config.js', 'BASE_URL 环境切换'],
        ['小程序', 'project.config.json', 'AppID、urlCheck'],
        ['公众号', 'application-local.yml wx.mp.*', 'AppID/Secret/Token'],
        ['公众号', '微信公众平台后台', '服务器 URL、IP 白名单、关联小程序'],
    ], [2.5, 5, 7.5])

    doc.add_paragraph()
    note = doc.add_paragraph(
        '安全提示：application-local.yml 含数据库密码、COS 密钥、微信 AppSecret，换机时请通过私密渠道拷贝，'
        '不要上传到公开 Git 仓库。若密钥曾泄露，请在腾讯云/微信后台轮换密钥。'
    )
    note.paragraph_format.first_line_indent = Cm(0.74)

    foot = doc.add_paragraph('—— 文档结束 ——')
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out))
    print(f'Generated: {out}')


if __name__ == '__main__':
    main()
