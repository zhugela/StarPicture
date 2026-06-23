"""
第二轮替换：补漏
1. "xxxx所测评中心" → "广州南方学院软工xxxx测评小组"
2. "年    月    日" → "2026 年 6 月 17 日"（计划/说明）"2026 年 6 月 19 日"（报告）
3. 模板说明 → 删
4. 在计划文档最前面填 4 人分工表（替换"小组项目分工及完成情况"）
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from pathlib import Path

BASE = Path("D:/code/StarPicture/docs/test")

# ============= 通用修复 =============
def fix_document(path, doc_date):
    doc = Document(path)

    # 修复"年    月    日"占位符
    for p in doc.paragraphs:
        if "年    月    日" in p.text:
            full = p.text.replace("年    月    日", doc_date)
            # 重新写：保留格式
            if p.runs:
                p.runs[0].text = full
                for r in p.runs[1:]:
                    r.text = ""
        if "xxxx所测评中心" in p.text:
            full = p.text.replace("xxxx所测评中心", "广州南方学院软工xxxx测评小组")
            if p.runs:
                p.runs[0].text = full
                for r in p.runs[1:]:
                    r.text = ""
        # 删模板说明
        if "【模板简要说明" in p.text or "中括号括住的内容" in p.text or "Cs001 是测试计划" in p.text or "并修改为正常字体和颜色】" in p.text or "蓝色文本处应替换" in p.text or "蓝色粗体处应替换" in p.text:
            # 整段删除
            p_elem = p._element
            p_elem.getparent().remove(p_elem)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "年    月    日" in p.text:
                        full = p.text.replace("年    月    日", doc_date)
                        if p.runs:
                            p.runs[0].text = full
                            for r in p.runs[1:]:
                                r.text = ""

    doc.save(path)
    print(f"已修复: {path}")


fix_document(BASE / "StarPicture_软件测试计划.docx", "2026 年 6 月 17 日")
fix_document(BASE / "StarPicture_软件测试说明.docx", "2026 年 6 月 17 日")
fix_document(BASE / "StarPicture_软件测试报告.docx", "2026 年 6 月 19 日")

# ============= 在计划文档前面填 4 人分工 =============
print("\n=== 在计划文档前面填 4 人分工 ===")
doc = Document(BASE / "StarPicture_软件测试计划.docx")

# 找到第一个段落（在标题"内娱图库（StarPicture）"之前）插入分工表
# 简单方法：在 body 最前面插段落+表格
body = doc.element.body
sectPr = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')

# 创建一个段落 + 表格，插到最前面
def add_to_top(doc, title, rows):
    """在最前面插入标题 + 表格"""
    # 新段落
    p = doc.paragraphs[0].insert_paragraph_before(title)
    p.style = doc.styles['Heading 2']

    # 新表格（不指定 style，用默认）
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    # 把表格移到最前
    tbl_elem = table._element
    body.remove(tbl_elem)
    # 找到第一个段落，把表格插到它前面
    first_p = doc.paragraphs[0]._element
    body.insert(list(body).index(first_p), tbl_elem)
    # 在表格后插入空行
    doc.paragraphs[0].insert_paragraph_before("")

add_to_top(doc, "小组项目分工及完成情况", [
    ["学号", "姓名", "专业", "主要工作", "成绩"],
    ["xxx01", "朱远亮（组长）", "计算机科学",
     "负责用户管理模块的测试；\n完成测试计划5.2、5.3小节撰写；\n参与测试用例第1，2章节的撰写；\n完成测试报告1、2章节撰写。",
     "待评"],
    ["xxx02", "李冠燃", "软件工程",
     "负责图片本地上传和关键字搜索模块的测试；\n完成测试计划5.4、5.5小节撰写；\n参与测试用例第3，4章节的撰写；\n完成测试报告3、4章节撰写。",
     "待评"],
    ["xxx03", "李坤纬", "软件工程",
     "负责空间创建和空间成员管理模块的测试；\n完成测试计划5.7、5.6小节撰写；\n参与测试用例第5章节的撰写；\n完成测试报告5章节撰写。",
     "待评"],
    ["xxx04", "林景彬", "软件工程",
     "负责文件本地上传和微信公众号门户的测试；\n完成测试计划5.7、5.6小节撰写；\n参与测试用例第6章节的撰写；\n完成测试报告6章节撰写。",
     "待评"],
])

add_to_top(doc, "测试用例、脚本工作量统计", [
    ["姓名", "测试类型", "功能用例", "性能用例", "接口用例", "安全用例", "合计"],
    ["朱远亮", "user 模块", 10, 1, 1, 1, 13],
    ["李冠燃", "picture 模块", 10, 1, 1, 1, 13],
    ["李坤纬", "space 模块", 10, 1, 1, 1, 13],
    ["林景彬", "file + wxMp 模块", 10, 1, 1, 1, 13],
    ["合计", "—", 40, 4, 4, 4, 52],
])

doc.save(BASE / "StarPicture_软件测试计划.docx")
print(f"已添加 4 人分工表到计划文档")

print("\n完成！")