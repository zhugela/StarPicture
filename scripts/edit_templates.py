"""
把 3 份模板的内容替换成 StarPicture 真实数据，保留模板原有的段落结构、标题层级、表格结构。
策略：
- 读取模板的 docx，遍历段落和表格
- 把所有可识别的占位符（如 "项目名称"、"编写"、"日期"、"XX系统"）替换成 StarPicture 内容
- 保留所有原有的字体/样式/格式
"""
import shutil
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

BASE = Path("D:/code/StarPicture/docs/test")
SRC_DIR = Path("D:/cxdownload")

REPLACEMENTS_GLOBAL = {
    # 项目信息
    "项目编号-CS001": "StarPicture-CS001",
    "项目编号-CS002": "StarPicture-CS002",
    "项目编号-CS003": "StarPicture-CS003",
    "FM-201202-0046": "StarPicture-2026-001",
    "XX 系统": "内娱图库（StarPicture）",
    "XX系统": "内娱图库（StarPicture）",
    "XX 系统测试": "StarPicture 系统测试",
    "项目名称": "内娱图库（StarPicture）",
    "软件测试计划": "软件测试计划（StarPicture）",
    # 小组信息
    "广州南方学院软工xxxx测评中心": "广州南方学院软工xxxx测评小组",
    "广州南方学院软工xxxx测评小组": "广州南方学院软工xxxx测评小组",
    "广州南方学院软工测评中心": "广州南方学院软工xxxx测评小组",
    # 日期
    "2014-XX-XX": "2026-06-17",
    "2021-XX-XX": "2026-06-17",
    # 文档号
    "CS001": "CS001",
    "CS002": "CS002",
    "CS003": "CS003",
    # 模板说明文字（一般是蓝色字提示删除）
    "【模板简要说明：在本文档中，凡蓝色文本处应替换为实际内容或删除， 并修改为正常字体和颜色】": "",
    "【模板简要说明：在本文档中，凡蓝色文本处应替换为实际内容或删除，并修改为正常字体和颜色】": "",
    "【模板简要说明：在本文档中，凡蓝色粗体处应替换为实际内容或删除， 并修改为正常字体和颜色】": "",
    "【模板简要说明：在本文档中，凡蓝色粗体处应替换为实际内容或删除，并修改为正常字体和颜色】": "",
    "【描述系统的业务背景、系统要实现的主要目标等信息。可从需求文档中直接将系统概述拷贝过来】": "内娱图库（StarPicture）是一款基于 Spring Boot 2.7.6 + MyBatis-Plus + MySQL 8 + 腾讯云 COS 构建的图片管理与分享平台，支持用户管理、图片上传/编辑/搜索、AI 扩图、空间协作、微信公众号接入等。",
    "【描述系统的业务背景、系统要实现的主要目标等信息】": "内娱图库（StarPicture）是一款基于 Spring Boot 2.7.6 + MyBatis-Plus + MySQL 8 + 腾讯云 COS 构建的图片管理与分享平台。",
    "【描述文档的编写目的】": "本文档是 StarPicture 项目测试活动的纲领性文件，明确本次测试的范围、策略、资源、进度与进入/通过/终止准则。",
    "【描述文档的编写目的。可从需求文档中直接将系统概述拷贝过来】": "本文档记录和说明 StarPicture 项目测试的过程及结果，为软件开发质量控制提供依据。",
}

REPLACEMENTS_REPORT_EXTRA = {
    "XX": "内娱图库（StarPicture）",
    "文档不仅用于说明和记录测试“XX系统”的过程及结果，还为“XX系统”的软件开发质量控制提供依据。": "本文档记录和说明 StarPicture 项目测试的过程及结果，为软件开发质量控制提供依据。",
}

def replace_in_doc(doc, replacements):
    """遍历 doc 中所有段落和表格，替换文字"""
    count = 0
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                # 替换 runs 中的文字（保留格式）
                full_text = para.text
                new_text = full_text.replace(old, new)
                if full_text != new_text:
                    # 清空所有 runs，保留第一个 run 的格式
                    if para.runs:
                        first_run = para.runs[0]
                        first_run.text = new_text
                        # 删除其他 runs
                        for run in para.runs[1:]:
                            run.text = ""
                        count += 1
                    else:
                        # 没有 run 直接加
                        para.add_run(new_text)
                        count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            full_text = para.text
                            new_text = full_text.replace(old, new)
                            if para.runs:
                                first_run = para.runs[0]
                                first_run.text = new_text
                                for run in para.runs[1:]:
                                    run.text = ""
                                count += 1
                            else:
                                para.add_run(new_text)
                                count += 1
    return count

# 处理 3 份文档
files = [
    ("软件测试计划_模板_20251027.docx", "StarPicture_软件测试计划.docx"),
    ("软件测试说明_模板_20250614.docx", "StarPicture_软件测试说明.docx"),
    ("软件测试报告_模板_20250618.docx", "StarPicture_软件测试报告.docx"),
]

for src_name, dst_name in files:
    src = SRC_DIR / src_name
    dst = BASE / dst_name

    doc = Document(src)
    # 替换全局占位符
    n1 = replace_in_doc(doc, REPLACEMENTS_GLOBAL)
    # 报告额外替换
    if "报告" in src_name:
        n2 = replace_in_doc(doc, REPLACEMENTS_REPORT_EXTRA)
        print(f"{dst_name}: 全局 {n1} 处 + 报告额外 {n2} 处")
    else:
        print(f"{dst_name}: 全局 {n1} 处")
    doc.save(dst)

print("\n完成！3 份模板已基于原模板改写为 StarPicture 内容。")